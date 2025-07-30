"""
Rory's AI Assistant 
===============================================

A comprehensive AI assistant system with GraphRAG technology for Rory Chen's professional profile.
Features: GraphRAG integration, OCR support, conversation tracking, and professional CV processing.

Author: Rory Chen
"""

import streamlit as st
import os
import sys
from sentence_transformers import SentenceTransformer
import numpy as np
from sklearn.metrics.pairwise import cosine_similarity
import glob
from typing import List, Dict, Tuple
import PyPDF2
import fitz 
from PIL import Image
import io
import sqlite3
from datetime import datetime
import uuid
import networkx as nx
from collections import defaultdict, Counter
import re
import requests
import json
from docx import Document

class GraphRAG:
    """Graph-based Retrieval Augmented Generation system"""
    
    def __init__(self):
        self.knowledge_graph = nx.Graph()
        self.entity_embeddings = {}
        self.document_entities = defaultdict(list)
        self.entity_documents = defaultdict(list)
        self.model = SentenceTransformer('all-MiniLM-L6-v2')
    
    def extract_entities(self, text: str, doc_id: str) -> List[str]:
        """Extract entities from text using enhanced pattern matching for better reasoning"""
        entities = []
        
        # Enhanced professional entities with more comprehensive patterns
        patterns = {
            'companies': r'\b(?:China CITIC Bank|CITIC Bank|AXA|Cigna|Ipsos|City University|Education University|Hong Kong|Macau|International)\b',
            'positions': r'\b(?:AVP|Assistant.*Manager|Data Science Analyst|Research Executive|Manager|Director|Senior|Lead|Analyst|Executive)\b',
            'skills': r'\b(?:Python|R|SQL|Machine Learning|Deep Learning|NLP|Tableau|Power BI|Azure|AWS|AutoML|MLOps|Computer Vision|Predictive Modeling|Analytics|Statistics|Data Science|AI|Artificial Intelligence)\b',
            'years': r'\b(?:20\d{2})\b',
            'degrees': r'\b(?:Master|Bachelor|MBA|PhD).*?(?:Arts|Science|Business|Engineering|Policy|Management|Quantitative|Analysis)\b',
            'industries': r'\b(?:Insurance|Banking|Healthcare|Finance|Market Research|Medicare|Retail|Technology)\b',
            'projects': r'\b(?:Dashboard|Framework|Pipeline|Model|System|Platform|Migration|Implementation)\b',
            'metrics': r'\b(?:\d+(?:\.\d+)?x|uplift|\d+%|\d+\+|monthly active users|coding effort|business impact)\b'
        }
        
        for category, pattern in patterns.items():
            matches = re.findall(pattern, text, re.IGNORECASE)
            for match in matches:
                entity = match.strip()
                if entity and len(entity) > 2:
                    entities.append(f"{category}:{entity}")
                    self.document_entities[doc_id].append(entity)
                    self.entity_documents[entity].append(doc_id)
        
        # Extract career progression relationships
        career_patterns = [
            r'(?:from|after|before)\s+([A-Z][a-zA-Z\s]+)\s+(?:to|at|joined)',
            r'(?:clients?|customers?)\s+(?:at|in|from)\s+([A-Z][a-zA-Z\s]+)',
            r'(?:experience|worked|served)\s+(?:at|in|with)\s+([A-Z][a-zA-Z\s]+)'
        ]
        
        for pattern in career_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            for match in matches:
                entity = match.strip()
                if entity and len(entity) > 2:
                    entities.append(f"career_connection:{entity}")
                    self.document_entities[doc_id].append(entity)
                    self.entity_documents[entity].append(doc_id)
        
        return entities
    
    def build_knowledge_graph(self, documents_dict: Dict[str, List[Tuple[str, Dict]]]):
        """Build knowledge graph from documents"""
        for doc_type, docs in documents_dict.items():
            for i, (text, metadata) in enumerate(docs):
                doc_id = f"{doc_type}_{i}_{metadata.get('source', 'unknown')}"
                
                # Extract entities from document
                entities = self.extract_entities(text, doc_id)
                
                # Add entities to graph
                for entity in entities:
                    if not self.knowledge_graph.has_node(entity):
                        self.knowledge_graph.add_node(entity, 
                                                    type=entity.split(':')[0] if ':' in entity else 'general',
                                                    documents=[doc_id])
                    else:
                        self.knowledge_graph.nodes[entity]['documents'].append(doc_id)
                
                # Create edges between co-occurring entities
                for j, entity1 in enumerate(entities):
                    for entity2 in entities[j+1:]:
                        if self.knowledge_graph.has_edge(entity1, entity2):
                            self.knowledge_graph[entity1][entity2]['weight'] += 1
                        else:
                            self.knowledge_graph.add_edge(entity1, entity2, weight=1, documents=[doc_id])
        
        # Generate embeddings for entities
        entity_names = list(self.knowledge_graph.nodes())
        if entity_names:
            embeddings = self.model.encode(entity_names, show_progress_bar=False)
            for entity, embedding in zip(entity_names, embeddings):
                self.entity_embeddings[entity] = embedding
    
    def graph_based_retrieval(self, query: str, k: int = 5) -> List[Tuple[str, float, List[str]]]:
        """Retrieve relevant information using graph structure"""
        if not self.knowledge_graph.nodes():
            return []
        
        # Find relevant entities based on query
        query_embedding = self.model.encode([query], show_progress_bar=False)[0]
        
        entity_scores = {}
        for entity, embedding in self.entity_embeddings.items():
            similarity = cosine_similarity([query_embedding], [embedding])[0][0]
            entity_scores[entity] = similarity
        
        # Get top entities
        top_entities = sorted(entity_scores.items(), key=lambda x: x[1], reverse=True)[:k]
        
        # Expand search using graph connections
        expanded_entities = set()
        for entity, score in top_entities:
            expanded_entities.add(entity)
            # Add connected entities with weighted scores
            for neighbor in self.knowledge_graph.neighbors(entity):
                edge_weight = self.knowledge_graph[entity][neighbor]['weight']
                expanded_entities.add(neighbor)
        
        # Collect relevant documents and sources
        relevant_docs = defaultdict(float)
        source_info = defaultdict(list)
        
        for entity in expanded_entities:
            if entity in self.knowledge_graph.nodes():
                docs = self.knowledge_graph.nodes[entity].get('documents', [])
                entity_score = entity_scores.get(entity, 0)
                
                for doc in docs:
                    relevant_docs[doc] += entity_score
                    source_info[doc].append(entity)
        
        # Return top documents with scores and sources
        top_docs = sorted(relevant_docs.items(), key=lambda x: x[1], reverse=True)[:k]
        
        results = []
        for doc_id, score in top_docs:
            sources = source_info[doc_id]
            results.append((doc_id, score, sources))
        
        return results

class ConversationDatabase:
    """Database manager for storing Q&A conversations"""
    
    def __init__(self, db_path: str = "conversation_history.db"):
        self.db_path = db_path
        self.init_database()
    
    def init_database(self):
        """Initialize the database with required tables"""
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS conversations (
                id TEXT PRIMARY KEY,
                session_id TEXT,
                timestamp DATETIME,
                user_question TEXT,
                ai_response TEXT,
                response_time_ms INTEGER
            )
        ''')
        
        conn.commit()
        conn.close()
    
    def save_conversation(self, session_id: str, question: str, response: str, response_time_ms: int = 0):
        """Save a Q&A pair to the database"""
        try:
            conn = sqlite3.connect(self.db_path)
            cursor = conn.cursor()
            
            conversation_id = str(uuid.uuid4())
            timestamp = datetime.now()
            
            cursor.execute('''
                INSERT INTO conversations (id, session_id, timestamp, user_question, ai_response, response_time_ms)
                VALUES (?, ?, ?, ?, ?, ?)
            ''', (conversation_id, session_id, timestamp, question, response, response_time_ms))
            
            conn.commit()
            conn.close()
            return True
        except Exception as e:
            print(f"Database error: {str(e)}")
            return False
    
    def get_conversation_count(self) -> int:
        """Get total number of conversations stored"""
        try:
            conn = sqlite3.connect(self.db_path)
            cursor = conn.cursor()
            cursor.execute("SELECT COUNT(*) FROM conversations")
            count = cursor.fetchone()[0]
            conn.close()
            return count
        except:
            return 0

class DocumentProcessor:
    """Document processor with OCR support"""
    
    def __init__(self):
        self.ocr_available = False
        # Try to import OCR dependencies silently
        try:
            import pytesseract
            pytesseract.get_tesseract_version()
            self.ocr_available = True
        except Exception as e:
            self.ocr_available = False
    
    def extract_text_from_pdf(self, pdf_path: str) -> str:
        """Extract text from PDF using available methods"""
        text = ""
        
        try:
            # Method 1: Try PyMuPDF first
            doc = fitz.open(pdf_path)
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                page_text = page.get_text()
                if page_text.strip():
                    text += page_text + "\n"
                elif self.ocr_available:
                    # If OCR is available and no text found, try OCR
                    try:
                        import pytesseract
                        pix = page.get_pixmap()
                        img_data = pix.tobytes("png")
                        img = Image.open(io.BytesIO(img_data))
                        ocr_text = pytesseract.image_to_string(img, lang='chi_sim+eng')
                        if ocr_text.strip():
                            text += f"[OCR Page {page_num + 1}]\n{ocr_text}\n"
                    except Exception as ocr_error:
                        st.warning(f"OCR failed for page {page_num + 1}: {str(ocr_error)}")
            doc.close()
            
            if len(text.strip()) > 100:
                return text
                
        except Exception as e:
            st.warning(f"PyMuPDF extraction failed: {str(e)}")
        
        try:
            # Method 2: Fallback to PyPDF2
            with open(pdf_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    page_text = page.extract_text()
                    if page_text.strip():
                        text += page_text + "\n"
            
            if len(text.strip()) > 100:
                return text
                
        except Exception as e:
            st.warning(f"PyPDF2 extraction failed: {str(e)}")
        
        if self.ocr_available:
            try:
                # Method 3: Full OCR as last resort
                import pytesseract
                st.info(f"Attempting OCR extraction for {os.path.basename(pdf_path)}...")
                doc = fitz.open(pdf_path)
                for page_num in range(min(10, len(doc))):
                    page = doc.load_page(page_num)
                    pix = page.get_pixmap()
                    img_data = pix.tobytes("png")
                    img = Image.open(io.BytesIO(img_data))
                    
                    ocr_text = pytesseract.image_to_string(img, lang='chi_sim+eng')
                    if ocr_text.strip():
                        text += f"[OCR Page {page_num + 1}]\n{ocr_text}\n"
                doc.close()
                
                if len(text.strip()) > 50:
                    return text
                    
            except Exception as e:
                st.error(f"OCR extraction failed: {str(e)}")
        
        return f"Failed to extract readable text from {os.path.basename(pdf_path)}."
    
    def extract_text_from_docx(self, docx_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            doc = Document(docx_path)
            text = ""
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += cell.text + "\n"
            
            if len(text.strip()) > 50:
                return text
            else:
                return f"No readable text found in {os.path.basename(docx_path)}"
                
        except Exception as e:
            st.error(f"Failed to extract text from DOCX {docx_path}: {str(e)}")
            return f"Failed to extract readable text from {os.path.basename(docx_path)}."
    
    def process_documents(self, directory: str) -> Dict[str, List[Tuple[str, Dict]]]:
        """Process documents with type separation"""
        documents = {
            "professional": [],
            "mindset": [],
            "general": []
        }
        
        # Process text files
        text_extensions = ['*.txt', '*.md']
        for ext in text_extensions:
            files = glob.glob(os.path.join(directory, ext))
            for file_path in files:
                try:
                    with open(file_path, 'r', encoding='utf-8') as f:
                        text = f.read()
                    
                    if text.strip() and len(text) > 100:
                        filename = os.path.basename(file_path).lower()
                        
                        if 'cv' in filename or 'resume' in filename:
                            doc_type = "professional"
                        elif 'flomo' in filename:
                            doc_type = "mindset"
                        else:
                            doc_type = "general"
                        
                        documents[doc_type].append((text, {
                            "source": os.path.basename(file_path), 
                            "type": "text",
                            "path": file_path,
                            "category": doc_type
                        }))
                        
                except Exception as e:
                    st.warning(f"Failed to read {file_path}: {str(e)}")
        
        # Process PDF files
        pdf_files = glob.glob(os.path.join(directory, '*.pdf'))
        for pdf_path in pdf_files:
            try:
                text = self.extract_text_from_pdf(pdf_path)
                if text.strip() and len(text) > 50:
                    filename = os.path.basename(pdf_path).lower()
                    
                    if 'cv' in filename or 'resume' in filename:
                        doc_type = "professional"
                    elif 'flomo' in filename:
                        doc_type = "mindset"
                    else:
                        doc_type = "general"
                    
                    documents[doc_type].append((text, {
                        "source": os.path.basename(pdf_path), 
                        "type": "pdf",
                        "path": pdf_path,
                        "category": doc_type
                    }))
            except Exception as e:
                pass  # Silent processing
        
        # Process DOCX files
        docx_files = glob.glob(os.path.join(directory, '*.docx'))
        for docx_path in docx_files:
            try:
                text = self.extract_text_from_docx(docx_path)
                if text.strip() and len(text) > 50:
                    filename = os.path.basename(docx_path).lower()
                    
                    if 'cv' in filename or 'resume' in filename or 'project' in filename or 'experience' in filename:
                        doc_type = "professional"
                    elif 'flomo' in filename:
                        doc_type = "mindset"
                    else:
                        doc_type = "general"
                    
                    documents[doc_type].append((text, {
                        "source": os.path.basename(docx_path), 
                        "type": "docx",
                        "path": docx_path,
                        "category": doc_type
                    }))
            except Exception as e:
                pass  # Silent processing
        
        return documents

class VectorStore:
    """Vector store with document type awareness"""
    
    def __init__(self):
        self.model = SentenceTransformer('all-MiniLM-L6-v2')
        self.documents = {
            "professional": [],
            "mindset": [],
            "general": []
        }
        self.embeddings = {
            "professional": [],
            "mindset": [],
            "general": []
        }
        self.metadata = {
            "professional": [],
            "mindset": [],
            "general": []
        }
    
    def add_documents_by_type(self, documents_dict: Dict[str, List[Tuple[str, Dict]]]):
        """Add documents categorized by type"""
        for doc_type, docs in documents_dict.items():
            for text, metadata in docs:
                if text.strip():
                    chunks = self._split_text(text, chunk_size=500, overlap=100)
                    for i, chunk in enumerate(chunks):
                        if chunk.strip() and len(chunk) > 50:
                            self.documents[doc_type].append(chunk)
                            chunk_metadata = metadata.copy()
                            chunk_metadata['chunk_index'] = i
                            self.metadata[doc_type].append(chunk_metadata)
        
        # Generate embeddings for each document type
        for doc_type in self.documents:
            if self.documents[doc_type]:
                self.embeddings[doc_type] = self.model.encode(
                    self.documents[doc_type], 
                    batch_size=32, 
                    show_progress_bar=False
                )
    
    def _split_text(self, text: str, chunk_size: int = 200, overlap: int = 20) -> List[str]:
        """Split text into semantic chunks based on content sections"""
        # First, try to split by clear section markers
        sections = []
        
        # Split by common CV/resume section markers
        section_markers = [
            r'\n\s*(?:PROFESSIONAL EXPERIENCE|WORK EXPERIENCE|EXPERIENCE|EMPLOYMENT)\s*\n',
            r'\n\s*(?:EDUCATION|ACADEMIC BACKGROUND)\s*\n',
            r'\n\s*(?:TECHNICAL SKILLS|SKILLS|COMPETENCIES)\s*\n',
            r'\n\s*(?:ACHIEVEMENTS|ACCOMPLISHMENTS|KEY ACHIEVEMENTS)\s*\n',
            r'\n\s*(?:PROJECTS|KEY PROJECTS)\s*\n',
            r'\n\s*(?:CONTACT|CONTACT INFORMATION)\s*\n',
            r'\n\s*(?:SUMMARY|PROFILE|OBJECTIVE)\s*\n'
        ]
        
        import re
        current_text = text
        
        # Try to identify and separate different content types
        for marker in section_markers:
            if re.search(marker, current_text, re.IGNORECASE):
                parts = re.split(marker, current_text, flags=re.IGNORECASE)
                if len(parts) > 1:
                    sections.extend(parts)
                    break
        
        if not sections:
            # If no clear sections, split by job positions or companies
            job_patterns = [
                r'\n\s*(?:AVP|Assistant.*Manager|Data Science Analyst|Research Executive|Manager|Director|Senior|Lead)\s+[^\n]*\n',
                r'\n\s*(?:China CITIC Bank|AXA|Cigna|Ipsos)[^\n]*\n'
            ]
            
            for pattern in job_patterns:
                if re.search(pattern, current_text, re.IGNORECASE):
                    parts = re.split(pattern, current_text, flags=re.IGNORECASE)
                    if len(parts) > 1:
                        sections.extend(parts)
                        break
        
        if not sections:
            # Fallback to paragraph-based splitting
            sections = text.split('\n\n')
        
        # Now chunk each section appropriately
        chunks = []
        for section in sections:
            section = section.strip()
            if not section or len(section) < 20:
                continue
            
            # If section is small enough, keep as one chunk
            if len(section.split()) <= chunk_size:
                chunks.append(section)
            else:
                # Split larger sections by sentences
                sentences = section.replace('\n', ' ').split('. ')
                current_chunk = []
                current_length = 0
                
                for sentence in sentences:
                    sentence = sentence.strip()
                    if not sentence:
                        continue
                        
                    sentence_length = len(sentence.split())
                    
                    if current_length + sentence_length > chunk_size and current_chunk:
                        chunks.append('. '.join(current_chunk) + '.')
                        # Smaller overlap for better separation
                        overlap_sentences = current_chunk[-1:] if current_chunk else []
                        current_chunk = overlap_sentences + [sentence]
                        current_length = sum(len(s.split()) for s in current_chunk)
                    else:
                        current_chunk.append(sentence)
                        current_length += sentence_length
                
                if current_chunk:
                    chunks.append('. '.join(current_chunk) + '.')
        
        return chunks if chunks else [text]
    
    def similarity_search_by_type(self, query: str, doc_type: str = "professional", k: int = 3) -> List[Tuple[str, Dict]]:
        """Search within specific document type"""
        if not self.documents[doc_type]:
            return []
        
        query_embedding = self.model.encode([query], show_progress_bar=False)
        similarities = cosine_similarity(query_embedding, self.embeddings[doc_type])[0]
        
        top_indices = np.argsort(similarities)[::-1][:k]
        
        results = []
        for idx in top_indices:
            if similarities[idx] > 0.2:  # Higher threshold for better quality
                results.append((self.documents[doc_type][idx], self.metadata[doc_type][idx]))
        
        return results

class GenAIResponseGenerator:
    """GenAI-powered response generator using Qwen Max API for intelligent reasoning"""
    
    def __init__(self, api_key=None):
        self.api_url = "https://dashscope.aliyuncs.com/api/v1/services/aigc/text-generation/generation"
        self.api_key = api_key or os.getenv("QWEN_API_KEY")
        self.fallback_available = True
        
        # Check if API key is available and test connection
        if not self.api_key:
            st.warning("⚠️ Qwen API key not found. Please enter your API key in the sidebar.")
            self.fallback_available = False
        else:
            # Test API connection
            try:
                test_response = self._test_api_connection()
                if test_response:
                    st.success("✅ Qwen Max API connected successfully!")
                else:
                    st.warning("⚠️ Qwen API connection failed. Using fallback system.")
                    self.fallback_available = False
            except Exception as e:
                st.warning(f"⚠️ Qwen API test failed: {str(e)}. Using fallback system.")
                self.fallback_available = False
    
    def _test_api_connection(self) -> bool:
        """Test API connection with a simple request"""
        try:
            headers = {
                "Authorization": f"Bearer {self.api_key}",
                "Content-Type": "application/json"
            }
            
            data = {
                "model": "qwen-max",
                "input": {
                    "prompt": "Hello, please respond with 'API connection successful'"
                },
                "parameters": {
                    "max_tokens": 50,
                    "temperature": 0.1
                }
            }
            
            response = requests.post(self.api_url, headers=headers, json=data, timeout=10)
            
            if response.status_code == 200:
                result = response.json()
                if 'output' in result and 'text' in result['output']:
                    return True
            
            return False
            
        except Exception as e:
            return False
    
    def generate_response(self, query: str, context: str, graph_sources: List = None, similarity_scores: List = None) -> str:
        """Generate response using Qwen Max API only - no fallback"""
        if not self.api_key or not self.fallback_available:
            return "❌ Qwen Max API is not available. Please configure your API key in the sidebar to use this service."
        
        try:
            return self._generate_with_api(query, context, graph_sources)
        except Exception as e:
            return f"❌ Qwen Max API error: {str(e)}. Please check your API key and try again."
    
    def _generate_with_api(self, query: str, context: str, graph_sources: List = None) -> str:
        """Generate response using Qwen Max API"""
        try:
            # Build GraphRAG context information
            graph_context = ""
            if graph_sources:
                graph_context = "\n\nGraphRAG Analysis:\n"
                for i, source_info in enumerate(graph_sources[:3]):
                    if isinstance(source_info, tuple) and len(source_info) >= 3:
                        doc_id, score, entities = source_info
                        graph_context += f"- Source {i+1}: {doc_id} (relevance: {score:.3f})\n"
                        graph_context += f"  Related entities: {', '.join(entities[:5])}\n"
            
            # Create the prompt for Qwen Max
            system_prompt = """You are Rory Chen's professional AI assistant with access to his comprehensive career information through GraphRAG technology. 

Your role is to provide intelligent, contextual responses about Rory's professional background, technical expertise, career achievements, and industry experience.

Key guidelines:
1. Use the provided context and GraphRAG analysis to give accurate, detailed responses
2. Demonstrate reasoning by connecting different pieces of information
3. Be professional and informative
4. If asked about career progression, explain the logical connections between roles and industries
5. Highlight technical achievements with specific metrics when available
6. For hiring/recruitment questions, provide comprehensive assessments
7. Always maintain a professional, confident tone
8. If information is insufficient, direct users to contact Rory directly: chengy823@gmail.com

Remember: You have access to advanced GraphRAG technology that can understand relationships between companies, skills, industries, and career progression."""

            user_prompt = f"""Question: {query}

Context from Rory's professional documents:
{context}

{graph_context}

Please provide a comprehensive, intelligent response that demonstrates reasoning and connects relevant information from the context. Show how different pieces of information relate to each other, especially for career progression and technical expertise questions."""

            # Combine system and user prompts for Qwen Max
            combined_prompt = f"{system_prompt}\n\n{user_prompt}"

            headers = {
                "Authorization": f"Bearer {self.api_key}",
                "Content-Type": "application/json"
            }
            
            data = {
                "model": "qwen-max",
                "input": {
                    "prompt": combined_prompt
                },
                "parameters": {
                    "max_tokens": 1000,
                    "temperature": 0.7,
                    "top_p": 0.8
                }
            }
            
            response = requests.post(self.api_url, headers=headers, json=data, timeout=30)
            
            if response.status_code == 200:
                result = response.json()
                ai_response = result['output']['text'].strip()
                
                # Add GraphRAG attribution
                if graph_sources:
                    ai_response += "\n\n*Response generated using advanced GraphRAG technology.*"
                
                return ai_response
            else:
                raise Exception(f"API request failed with status {response.status_code}: {response.text}")
                
        except Exception as e:
            raise Exception(f"API generation error: {str(e)}")
    
    def _generate_with_fallback(self, query: str, context: str, graph_sources: List = None) -> str:
        """Fallback response generation using intelligent context analysis"""
        try:
            # Intelligent context-based response generation
            query_lower = query.lower()
            
            # Extract key information from context
            context_lines = context.split('\n')
            key_info = []
            
            # Look for specific information patterns
            for line in context_lines:
                line = line.strip()
                if any(keyword in line.lower() for keyword in ['avp', 'manager', 'analyst', 'executive']):
                    key_info.append(line)
                elif any(keyword in line.lower() for keyword in ['python', 'sql', 'machine learning', 'deep learning']):
                    key_info.append(line)
                elif any(keyword in line.lower() for keyword in ['citic', 'axa', 'cigna', 'ipsos']):
                    key_info.append(line)
                elif any(keyword in line.lower() for keyword in ['master', 'bachelor', 'university']):
                    key_info.append(line)
            
            # Build intelligent response based on available context
            response_parts = []
            
            if key_info:
                response_parts.append("Based on Rory's professional background:")
                response_parts.append("")
                
                # Add relevant information
                for info in key_info[:8]:  # Limit to most relevant info
                    if info and len(info) > 10:
                        response_parts.append(f"• {info}")
                
                response_parts.append("")
                response_parts.append("Rory has demonstrated consistent career growth and technical expertise across multiple industries including market research, healthcare, insurance, and banking.")
            else:
                response_parts.append("Based on the available information about Rory Chen:")
                response_parts.append("")
                response_parts.append("Rory is an experienced data science professional with 8 years of career progression from Research Executive to AVP level, demonstrating strong technical and leadership capabilities.")
            
            # Add contact information for detailed inquiries
            response_parts.append("")
            response_parts.append("For more detailed information, please contact Rory directly:")
            response_parts.append("📧 Email: chengy823@gmail.com")
            response_parts.append("📍 Location: Hong Kong SAR")
            
            response = "\n".join(response_parts)
            
            # Add GraphRAG attribution
            if graph_sources:
                response += "\n\n*Response generated using GraphRAG technology.*"
            
            return response
            
        except Exception as e:
            return f"I encountered an error while processing your question: {str(e)}. Please contact Rory directly at chengy823@gmail.com for detailed information."
    
    def _generate_technical_skills_response(self, context: str) -> str:
        """Generate technical skills focused response"""
        return """Rory has developed comprehensive technical expertise throughout his 8-year career in data science:

**Programming Languages:**
- **Python**: Advanced proficiency for machine learning, data analysis, and automation
- **R**: Statistical computing, advanced analytics, and data visualization
- **SQL**: Database management, complex queries, and data extraction

**Machine Learning & AI:**
- Machine Learning frameworks and model development
- Deep Learning for complex pattern recognition
- Natural Language Processing (NLP) for text analytics
- Computer Vision for image and video analysis
- AutoML pipeline development and MLOps implementation

**Analytics & Visualization:**
- Predictive modeling and statistical analysis
- Business Intelligence and dashboard development
- Tableau and Power BI for data visualization
- Advanced analytics for business insights

**Cloud & Infrastructure:**
- Azure and AWS cloud platforms
- Databricks for big data processing
- Docker and Kubernetes for deployment
- Spark and Hadoop for distributed computing

**Industry Applications:**
- Healthcare analytics and claims processing
- Insurance modeling and risk assessment
- Banking and financial services analytics
- Market research and consumer behavior analysis"""
    
    def _generate_hiring_assessment_response(self, context: str) -> str:
        """Generate hiring assessment response"""
        return """Based on Rory's professional background, here's my assessment for GenAI management roles:

**Leadership & Management Experience:**
- **Current Role**: AVP of Data Science at China CITIC Bank International, leading data science initiatives
- **Career Progression**: 8 years of growth from Research Executive to AVP level (2017-2025)
- **Team Leadership**: Experience managing data science teams and cross-functional projects
- **Strategic Vision**: Led AI governance frameworks and enterprise-scale AI deployment

**Technical Expertise for GenAI:**
- **AI/ML Foundation**: 20+ ML models developed with measurable business impact (1.5x uplift)
- **Advanced Technologies**: Deep Learning, NLP, Computer Vision expertise
- **Automation**: Created AutoML pipeline reducing coding effort by 80%
- **Innovation**: Implemented Next-Best-Action frameworks and AI+BI solutions

**Business Impact & Results:**
- **Quantifiable Success**: 1.5x business uplift from ML models vs control groups
- **Scale**: Built systems with 100+ monthly active users
- **Efficiency**: 80% reduction in coding effort through automation
- **Cross-Industry**: Experience in healthcare, insurance, banking, and market research

**GenAI Management Readiness:**
- **Technical Depth**: Strong foundation in AI/ML technologies essential for GenAI
- **Leadership Experience**: Proven ability to lead technical teams and initiatives
- **Business Acumen**: Track record of delivering measurable business value
- **Innovation Mindset**: Experience with cutting-edge AI implementations

**Contact Information:**
📧 Email: chengy823@gmail.com
📍 Location: Hong Kong SAR

**Recommendation**: Rory would be an excellent fit for GenAI management roles, bringing both deep technical expertise and proven leadership capabilities."""
    
    def _generate_age_response(self, context: str) -> str:
        """Generate age calculation response"""
        return """Based on Rory's educational and career timeline, here's the age calculation:

**Educational Timeline:**
- **Bachelor's Degree**: Education University of Hong Kong (2011-2014)
- **Master's in Public Policy**: City University of Hong Kong (2014-2015)
- **Master's in Quantitative Analysis**: City University of Hong Kong (2016-2017)

**Career Timeline (8 years: 2017-2025):**
- **Research Executive** at Ipsos Hong Kong (2017-2018)
- **Data Science Analyst** at Cigna International (2018-2019)
- **Assistant Data Science Manager** at AXA (2019-2022)
- **AVP of Data Science** at China CITIC Bank (2022-Present)

**Age Estimation:**
Assuming Rory started university at the typical age of 18-19 in 2011:
- **Birth Year**: Approximately 1992-1993
- **Current Age (2025)**: Approximately **32-33 years old**

This age aligns well with his career progression, showing 8 years of professional experience and advancement to AVP level, which is typical for someone in their early 30s with strong performance and leadership capabilities."""
    
    def _generate_achievements_response(self, context: str) -> str:
        """Generate achievements focused response"""
        return """Rory's key achievements in analytics and data science include:

**🏆 Technical Achievements:**
- **20+ ML Models Developed**: Achieved measurable business impact with 1.5x business uplift vs control groups
- **AutoML Pipeline Creation**: Reduced coding effort by 80% for faster model development and deployment
- **Next-Best-Action (NBA) Frameworks**: Implemented intelligent decision-making systems for business optimization
- **Cloud Migration Leadership**: Successfully led migration projects from on-premise to Azure infrastructure

**📈 Business Impact:**
- **AI+BI Framework Design**: Created comprehensive framework with ONE AXA Dashboard serving 100+ monthly active users
- **Insurance Analytics Advancement**: Enhanced data capabilities for predictive modeling and risk assessment
- **Retail Banking AI**: Led data science initiatives with AI-driven decision-making systems
- **US Medicare Analytics**: Supported predictive modeling and cost forecasting for healthcare analytics

**🚀 Leadership & Innovation:**
- **Career Progression**: 8-year advancement from Research Executive to AVP level demonstrating leadership growth
- **Cross-Industry Expertise**: Successfully applied data science across market research, healthcare, insurance, and banking
- **AI Governance**: Implemented enterprise-scale AI model governance frameworks
- **Team Development**: Led data science teams and cross-functional initiatives

**🎯 Industry Recognition:**
- **Proven ROI**: Consistent delivery of production AI systems with measurable business value
- **Technical Leadership**: Recognition through progressive career advancement and increased responsibilities
- **Innovation Driver**: Pioneer in AutoML and AI+BI integration within organizations
- **Mentorship**: Experience developing data science capabilities and teams across multiple industries

**📊 Quantifiable Results:**
- 1.5x business performance improvement
- 80% efficiency gain in model development
- 100+ monthly active users for developed systems
- 20+ successful ML model deployments
- 8 years of consistent career growth"""
    
    def _generate_experience_response(self, context: str) -> str:
        """Generate experience focused response"""
        return """Rory has built an impressive 8-year career in data science across multiple industries:

**Career Progression (2017-2025):**

**🏦 China CITIC Bank International** - AVP, Data Science (Nov 2022 - Current)
- Leading data science initiatives in retail banking with AI-driven decision-making
- Developing and deploying ML models for business optimization
- Managing data science teams and strategic AI implementations

**🛡️ AXA Hong Kong and Macau** - Assistant Data Science Manager (Sep 2019 - Nov 2022)
- Advanced data capabilities for insurance analytics and predictive modeling
- Designed AI+BI framework and ONE AXA Dashboard with 100+ monthly users
- Led cloud migration from on-premise to Azure infrastructure
- Implemented enterprise-scale data science solutions

**🏥 Cigna International Office** - Data Science Analyst (Aug 2018 - Sep 2019)
- Supported US Medicare analytics with focus on predictive modeling
- Developed healthcare cost forecasting models
- Analyzed claims data for business insights

**📊 Ipsos Hong Kong** - Research Executive (2017 - 2018)
- Conducted market research and data analysis for various industries
- Statistical modeling for consumer behavior analysis
- Foundation in research methodology and data analysis

**Key Highlights:**
- **Industry Diversity**: Experience across market research, healthcare, insurance, and banking sectors
- **Progressive Growth**: Consistent advancement from analyst to AVP level in 8 years
- **Technical Leadership**: Evolution from individual contributor to team leader and strategic decision maker
- **Business Impact**: Proven track record of delivering measurable results across different business contexts
- **Cross-Functional Skills**: Combination of technical expertise and business acumen
- **Innovation Focus**: Consistent implementation of cutting-edge technologies and methodologies"""
    
    def _generate_contact_response(self, context: str) -> str:
        """Generate contact information response"""
        return """You can reach out to Rory through the following channels:

**📧 Primary Contact:**
Email: chengy823@gmail.com

**📍 Location:**
Hong Kong SAR

**🔗 Professional Network:**
LinkedIn: Available with 688 followers and 500+ connections

**💼 Current Position:**
AVP of Data Science at China CITIC Bank International

**🤝 Open to:**
- Collaboration opportunities
- Data science consulting
- Speaking engagements
- Professional networking
- Career discussions and mentorship

Rory is currently serving as AVP of Data Science and is open to discussing opportunities, collaborations, or sharing insights about data science and analytics."""
    
    def _generate_general_response(self, query: str, context: str) -> str:
        """Generate general response using available context"""
        # Extract key information from context
        context_summary = context[:800] + "..." if len(context) > 800 else context
        
        return f"""Based on the available information about Rory Chen:

{context_summary}

Rory has 8 years of professional experience (2017-2025) and has progressed from Research Executive to AVP level, demonstrating strong technical and leadership capabilities across multiple industries including market research, healthcare, insurance, and banking.

For more specific information, please feel free to ask about:
- Technical skills and programming expertise
- Career experience and achievements
- Industry experience and projects
- Contact information
- Educational background
- Leadership and management experience"""

class ReasoningChain:
    """Main reasoning system with GraphRAG and GenAI response generation"""
    
    def __init__(self, vector_store: VectorStore, graph_rag: GraphRAG = None):
        self.vector_store = vector_store
        self.graph_rag = graph_rag
        self.response_generator = GenAIResponseGenerator(api_key=os.getenv("QWEN_API_KEY"))
    
    def __call__(self, inputs: Dict) -> Dict:
        """Process query with GraphRAG and generate response using Qwen API for ALL questions"""
        query = inputs.get("query", "").strip()
        
        # Use GraphRAG for enhanced retrieval if available
        graph_sources = []
        if self.graph_rag:
            graph_results = self.graph_rag.graph_based_retrieval(query, k=5)
            if graph_results:
                graph_sources = [(doc_id, score, sources) for doc_id, score, sources in graph_results]
        
        # Get professional context
        professional_docs = self.vector_store.similarity_search_by_type(query, "professional", k=3)
        
        # Build context from available information (even if empty)
        professional_context = "\n\n".join([doc[0] for doc in professional_docs]) if professional_docs else ""
        
        # ALWAYS use Qwen API for response generation, regardless of context availability
        response = self.response_generator.generate_response(
            query, 
            professional_context, 
            graph_sources
        )
        
        return {"result": response}

@st.cache_resource
def get_conversation_db():
    return ConversationDatabase()

@st.cache_resource
def initialize_ai_system():
    """Initialize the complete AI system"""
    try:
        current_dir = os.path.dirname(os.path.abspath(__file__))
        
        # Initialize document processor
        doc_processor = DocumentProcessor()
        
        # Process documents
        documents_dict = doc_processor.process_documents(current_dir)
        
        # Fallback content if no documents found
        if not any(documents_dict.values()):
            fallback_content = """
            Rory Chen - Data Science & Analytics Expert
            
            Professional Experience (8 years: 2017-2025):
            AVP, Data Science at China CITIC Bank International (Nov 2022 - Current)
            - Led data science initiatives in retail banking with AI-driven decision-making
            - Developed 20+ ML models achieving 1.5x business uplift vs control group
            - Created AutoML pipeline reducing coding effort by 80%
            - Implemented Next-Best-Action (NBA) model framework
            
            Assistant Data Science Manager at AXA Hong Kong and Macau (Sep 2019 - Nov 2022)
            - Advanced data capabilities for insurance analytics and predictive modeling
            - Designed AI+BI framework and ONE AXA Dashboard (100+ monthly users)
            - Led cloud migration from on-premise to Azure infrastructure
            
            Data Science Analyst at Cigna International Office (Aug 2018 - Sep 2019)
            - Supported US Medicare analytics with predictive modeling focus
            
            Research Executive at Ipsos Hong Kong (2017 - 2018)
            - Conducted market research and data analysis for various industries
            
            Education:
            Master of Arts in Quantitative Analysis for Business, City University of Hong Kong (2016-2017)
            Master of Arts in Public Policy and Management, City University of Hong Kong (2014-2015)
            Bachelor of Social Sciences (Honours), Education University of Hong Kong (2011-2014)
            
            Technical Skills:
            Programming: Python, R, SQL (Advanced proficiency)
            Machine Learning: Deep Learning, NLP, Computer Vision, MLOps
            Analytics: Predictive Modeling, Customer Analytics, Time Series Analysis
            Cloud Platforms: Azure, AWS, Databricks, Cloudera CDSW
            Visualization: Tableau, Power BI, Dashboard Development
            Tools: Spark, Hadoop, Docker, Kubernetes, AutoML
            
            Contact: chengy823@gmail.com
            Location: Hong Kong SAR
            """
            documents_dict["professional"] = [(fallback_content, {"source": "profile", "type": "text", "category": "professional"})]
        
        # Create vector store
        vector_store = VectorStore()
        vector_store.add_documents_by_type(documents_dict)
        
        # Initialize GraphRAG system
        graph_rag = GraphRAG()
        graph_rag.build_knowledge_graph(documents_dict)
        
        # Create reasoning chain
        reasoning_chain = ReasoningChain(vector_store, graph_rag)
        
        # Calculate totals
        total_docs = sum(len(docs) for docs in documents_dict.values())
        total_chunks = sum(len(chunks) for chunks in vector_store.documents.values())
        
        return reasoning_chain, total_docs, total_chunks
    
    except Exception as e:
        st.error(f"Failed to initialize AI system: {str(e)}")
        return None, 0, 0

# Page configuration
st.set_page_config(
    page_title="Rory's AI Assistant",
    page_icon="🤖",
    layout="wide",
    initial_sidebar_state="collapsed"
)

# Custom CSS for professional styling
st.markdown("""
<style>
    .main .block-container {
        font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
        font-size: 16px;
        line-height: 1.6;
    }
    
    .main-header {
        text-align: center;
        padding: 2rem 0;
        background: linear-gradient(90deg, #1e3c72 0%, #2a5298 100%);
        color: white;
        border-radius: 10px;
        margin-bottom: 2rem;
    }
    
    .main-header h1 {
        font-size: 2.5rem !important;
        font-weight: 700 !important;
        margin-bottom: 0.5rem !important;
    }
    
    .main-header h3 {
        font-size: 1.5rem !important;
        font-weight: 400 !important;
        margin-bottom: 1rem !important;
    }
    
    .main-header p {
        font-size: 1.1rem !important;
        opacity: 0.9;
    }
    
    .status-info {
        background: #e8f5e8;
        padding: 1rem;
        border-radius: 8px;
        border-left: 4px solid #4caf50;
        margin-bottom: 1rem;
        font-size: 1rem;
    }
    
    .chat-container {
        background: white;
        padding: 1.5rem;
        border-radius: 10px;
        box-shadow: 0 2px 10px rgba(0,0,0,0.1);
    }
    
    .stButton > button {
        width: 100%;
        border-radius: 8px;
        border: 1px solid #ddd;
        background: white;
        color: #333;
        padding: 0.75rem 1rem;
        margin: 0.3rem 0;
        transition: all 0.3s;
        font-size: 0.95rem;
        font-weight: 500;
        text-align: left;
    }
    
    .stButton > button:hover {
        background: #f0f8ff;
        border-color: #3498db;
        transform: translateY(-1px);
        box-shadow: 0 2px 8px rgba(52, 152, 219, 0.2);
    }
    
    .stChatMessage {
        font-size: 1rem !important;
        line-height: 1.6 !important;
    }
    
    .stChatMessage [data-testid="chatAvatarIcon-user"] {
        background-color: #3498db !important;
    }
    
    .stChatMessage [data-testid="chatAvatarIcon-assistant"] {
        background-color: #2ecc71 !important;
    }
    
    .stTextInput > div > div > input {
        font-size: 1rem !important;
        padding: 0.75rem !important;
    }
</style>
""", unsafe_allow_html=True)

# Sidebar for API key configuration
with st.sidebar:
    st.markdown("### 🔑 API Configuration")
    qwen_api_key = st.text_input(
        "Qwen API Key:",
        value="sk-015ea57c8b254c4181d30b2de4259d8b",
        type="password",
        help="Enter your Qwen API key for enhanced AI responses"
    )
    
    if qwen_api_key:
        os.environ["QWEN_API_KEY"] = qwen_api_key
        st.success("✅ API key configured!")
    
    st.markdown("---")
    st.markdown("### 📊 System Info")
    st.markdown(f"**Documents**: {total_docs if 'total_docs' in locals() else 0}")
    st.markdown(f"**Conversations**: {conversation_db.get_conversation_count() if 'conversation_db' in locals() else 0}")

# Main header
st.markdown("""
<div class="main-header">
    <h1>🤖 Rory's AI Assistant</h1>
    <p>Advanced GraphRAG Technology | Professional Data Science Expert</p>
</div>
""", unsafe_allow_html=True)

# Initialize AI system and database
with st.spinner("🚀 Initializing AI System..."):
    qa_chain, total_docs, total_chunks = initialize_ai_system()
    conversation_db = get_conversation_db()

if qa_chain is None:
    st.error("❌ Rory's AI system is not available. Please check the system configuration.")
    st.stop()

# Initialize session ID for conversation tracking
if "session_id" not in st.session_state:
    st.session_state.session_id = str(uuid.uuid4())

# Get conversation count for display
total_conversations = conversation_db.get_conversation_count()

# Display system status
st.markdown(f"""
<div class="status-info">
    <strong style="color: #2c3e50;">🧠 Final AI System Status:</strong> <span style="color: #34495e;">{total_docs} documents processed, 
    {total_chunks} knowledge chunks available | 🎯 GraphRAG Enhanced | 💾 {total_conversations} conversations stored</span>
</div>
""", unsafe_allow_html=True)

# Create two columns for layout
col1, col2 = st.columns([1, 1])

with col1:
    # CV Section
    st.markdown("""
    <div style="background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); 
                padding: 2rem; border-radius: 15px; color: white; margin-bottom: 2rem;
                box-shadow: 0 8px 32px rgba(0,0,0,0.1);">
        <div style="text-align: center; margin-bottom: 1.5rem;">
            <h1 style="color: white; margin-bottom: 0.5rem;">👨‍💼 Rory Chen</h1>
            <h3 style="color: #e8f4fd; margin-bottom: 0.5rem;">Data Science & Analytics Expert</h3>
            <p style="color: #b8d4f0; font-size: 1.1rem;">🏢 China CITIC Bank International Limited</p>
            <p style="color: #b8d4f0;">📍 Hong Kong SAR | 🔗 LinkedIn: 688 followers, 500+ connections</p>
        </div>
    </div>
    
    <div style="background: white; padding: 2rem; border-radius: 15px; 
                box-shadow: 0 4px 20px rgba(0,0,0,0.08); margin-bottom: 2rem;">
        <h2 style="color: #2c3e50; border-bottom: 3px solid #e74c3c; padding-bottom: 0.5rem;">💼 Professional Experience (8 years: 2017-2025)</h2>
        <div style="margin: 1rem 0;">
            <div style="margin-bottom: 1.5rem;">
                <h4 style="color: #34495e; margin-bottom: 0.3rem;">🏦 China CITIC Bank International</h4>
                <p style="color: #7f8c8d; margin-bottom: 0.3rem;"><strong>AVP, Data Science</strong> (Nov 2022 - Current)</p>
                <p style="color: #7f8c8d; font-size: 0.9rem;">Led data science initiatives, developed 20+ ML models with 1.5x business uplift.</p>
            </div>
            <div style="margin-bottom: 1.5rem;">
                <h4 style="color: #34495e; margin-bottom: 0.3rem;">🛡️ AXA Hong Kong and Macau</h4>
                <p style="color: #7f8c8d; margin-bottom: 0.3rem;"><strong>Assistant Data Science Manager</strong> (Sep 2019 - Nov 2022)</p>
                <p style="color: #7f8c8d; font-size: 0.9rem;">Advanced data capabilities for insurance analytics and cloud migration.</p>
            </div>
            <div style="margin-bottom: 1.5rem;">
                <h4 style="color: #34495e; margin-bottom: 0.3rem;">🏥 Cigna International Office</h4>
                <p style="color: #7f8c8d; margin-bottom: 0.3rem;"><strong>Data Science Analyst</strong> (Aug 2018 - Sep 2019)</p>
                <p style="color: #7f8c8d; font-size: 0.9rem;">Supported US Medicare analytics with predictive modeling focus.</p>
            </div>
            <div style="margin-bottom: 1.5rem;">
                <h4 style="color: #34495e; margin-bottom: 0.3rem;">📊 Ipsos Hong Kong</h4>
                <p style="color: #7f8c8d; margin-bottom: 0.3rem;"><strong>Research Executive</strong> (2017 - 2018)</p>
                <p style="color: #7f8c8d; font-size: 0.9rem;">Conducted market research and statistical modeling for consumer behavior.</p>
            </div>
        </div>
    </div>
    
    <div style="background: white; padding: 2rem; border-radius: 15px; 
                box-shadow: 0 4px 20px rgba(0,0,0,0.08); margin-bottom: 2rem;">
        <h2 style="color: #2c3e50; border-bottom: 3px solid #3498db; padding-bottom: 0.5rem;">🎓 Education</h2>
        <div style="margin: 1rem 0;">
            <div style="margin-bottom: 1rem;">
                <h4 style="color: #34495e; margin-bottom: 0.3rem;">🏛️ City University of Hong Kong</h4>
                <p style="color: #7f8c8d; margin-bottom: 0.2rem;"><strong>Master of Arts in Quantitative Analysis for Business</strong> (2016-2017)</p>
                <p style="color: #7f8c8d; margin-bottom: 0.2rem;"><strong>Master of Arts in Public Policy and Management</strong> (2014-2015)</p>
            </div>
            <div>
                <h4 style="color: #34495e; margin-bottom: 0.3rem;">🏛️ Education University of Hong Kong</h4>
                <p style="color: #7f8c8d;"><strong>Bachelor of Social Sciences (Honours)</strong> (2011-2014)</p>
            </div>
        </div>
    </div>
    
    <div style="background: white; padding: 2rem; border-radius: 15px; 
                box-shadow: 0 4px 20px rgba(0,0,0,0.08);">
        <h2 style="color: #2c3e50; border-bottom: 3px solid #f39c12; padding-bottom: 0.5rem;">🚀 Core Expertise</h2>
        <div style="display: grid; grid-template-columns: 1fr 1fr; gap: 1rem; margin-top: 1rem;">
            <div>
                <h5 style="color: #2980b9; margin-bottom: 0.5rem;">📊 Data Science</h5>
                <p style="color: #7f8c8d; font-size: 0.9rem;">Python, R, SQL, Machine Learning</p>
            </div>
            <div>
                <h5 style="color: #27ae60; margin-bottom: 0.5rem;">🤖 AI & ML</h5>
                <p style="color: #7f8c8d; font-size: 0.9rem;">Deep Learning, NLP, Computer Vision</p>
            </div>
            <div>
                <h5 style="color: #8e44ad; margin-bottom: 0.5rem;">📈 Analytics</h5>
                <p style="color: #7f8c8d; font-size: 0.9rem;">Predictive Modeling, Business Intelligence</p>
            </div>
            <div>
                <h5 style="color: #e67e22; margin-bottom: 0.5rem;">🛠️ Tools</h5>
                <p style="color: #7f8c8d; font-size: 0.9rem;">Tableau, Power BI, Cloud Platforms</p>
            </div>
            <div>
                <h5 style="color: #c0392b; margin-bottom: 0.5rem;">🧠 GraphRAG</h5>
                <p style="color: #7f8c8d; font-size: 0.9rem;">Knowledge Graphs, Entity Extraction</p>
            </div>
            <div>
                <h5 style="color: #16a085; margin-bottom: 0.5rem;">🔗 LangChain</h5>
                <p style="color: #7f8c8d; font-size: 0.9rem;">RAG Systems, Vector Databases</p>
            </div>
            <div>
                <h5 style="color: #9b59b6; margin-bottom: 0.5rem;">🚀 Streamlit</h5>
                <p style="color: #7f8c8d; font-size: 0.9rem;">Interactive AI Applications</p>
            </div>
            <div>
                <h5 style="color: #f39c12; margin-bottom: 0.5rem;">📄 OCR & NLP</h5>
                <p style="color: #7f8c8d; font-size: 0.9rem;">Document Processing, Text Analytics</p>
            </div>
        </div>
    </div>
    """, unsafe_allow_html=True)

with col2:
    # Q&A Input Box Section
    st.markdown("### ❓ Ask Rory's AI Assistant")
    st.markdown("Type your question below or click on the sample questions:")
    
    # Direct Q&A input box
    user_question = st.text_input(
        "Your Question:",
        placeholder="Ask me anything about Rory's experience, skills, or expertise...",
        key="direct_question_input"
    )
    
    if st.button("🚀 Ask Question", key="ask_direct_question") and user_question.strip():
        st.session_state.selected_question = user_question.strip()
        st.rerun()
    
    st.markdown("---")
    
    # GraphRAG Showcase Section
    st.markdown("### 🧠 GraphRAG-Powered Intelligence")
    st.markdown("*Ask any question - the AI will use advanced GraphRAG technology to provide intelligent, contextual responses based on Rory's professional knowledge base.*")
    
    st.markdown("**Examples of GraphRAG reasoning:**")
    st.markdown("- *Why did Rory transition from Ipsos to AXA?*")
    st.markdown("- *How do Rory's skills in insurance analytics apply to banking?*") 
    st.markdown("- *What's the connection between Rory's clients and career progression?*")
    
    # Chat Interface
    st.markdown("---")
    st.markdown('<div class="chat-container">', unsafe_allow_html=True)
    
    # Chat header with clear button
    col_chat1, col_chat2 = st.columns([3, 1])
    with col_chat1:
        st.markdown("### 💬 Chat with Rory's AI")
    with col_chat2:
        if st.button("🗑️ Clear History", key="clear_history", help="Clear all conversation history"):
            st.session_state.messages = []
            st.session_state.session_id = str(uuid.uuid4())
            st.rerun()
    
    # Initialize chat history
    if "messages" not in st.session_state:
        st.session_state.messages = []
        # Add welcome message
        st.session_state.messages.append({
            "role": "assistant", 
            "content": "👋 Hello! I'm Rory's AI assistant powered by advanced GraphRAG technology. I can provide comprehensive information about Rory's professional background, technical expertise, and career achievements. Feel free to ask me anything!"
        })

    # Handle selected question from sample questions
    if "selected_question" in st.session_state:
        selected_q = st.session_state.selected_question
        st.session_state.messages.append({"role": "user", "content": selected_q})
        
        # Show processing indicator
        with st.spinner("🧠 Processing with GraphRAG..."):
            try:
                result = qa_chain({"query": selected_q})
                
                if isinstance(result, dict):
                    response = result.get("result", "Sorry, I couldn't find a relevant answer.")
                else:
                    response = str(result)
                
                st.session_state.messages.append({"role": "assistant", "content": response})
                
                # Save conversation to database
                conversation_db.save_conversation(
                    session_id=st.session_state.session_id,
                    question=selected_q,
                    response=response
                )
                
            except Exception as e:
                error_msg = f"Sorry, I encountered an error while processing your question: {str(e)}"
                st.session_state.messages.append({"role": "assistant", "content": error_msg})
        
        del st.session_state.selected_question
        st.rerun()

    # Display chat messages
    for message in st.session_state.messages:
        with st.chat_message(message["role"]):
            st.markdown(message["content"])

    # Chat input
    if prompt := st.chat_input("Ask me anything about Rory's experience and expertise..."):
        # Add user message to chat history
        st.session_state.messages.append({"role": "user", "content": prompt})
        
        # Display user message
        with st.chat_message("user"):
            st.markdown(prompt)
        
        # Generate assistant response
        with st.chat_message("assistant"):
            with st.spinner("🔍 Analyzing with GraphRAG..."):
                try:
                    result = qa_chain({"query": prompt})
                    
                    if isinstance(result, dict):
                        response = result.get("result", "Sorry, I couldn't find a relevant answer.")
                    else:
                        response = str(result)
                    
                    st.markdown(response)
                    st.session_state.messages.append({"role": "assistant", "content": response})
                    
                    # Save conversation to database
                    conversation_db.save_conversation(
                        session_id=st.session_state.session_id,
                        question=prompt,
                        response=response
                    )
                    
                except Exception as e:
                    error_msg = f"Sorry, I encountered an error while processing your question: {str(e)}"
                    st.error(error_msg)
                    st.session_state.messages.append({"role": "assistant", "content": error_msg})

    st.markdown('</div>', unsafe_allow_html=True)

# Footer
st.markdown("---")
st.markdown("""
<div style="text-align: center; padding: 2rem; color: #666;">
    <p>🤖 <strong>Rory's AI Assistant - Final Consolidated Version</strong> | Powered by GraphRAG Technology | 📄 PDF Processing with OCR</p>
    <p>📧 For direct contact: chengy823@gmail.com</p>
    <p><em>Advanced AI system with GraphRAG technology for comprehensive professional insights</em></p>
</div>
""", unsafe_allow_html=True)
