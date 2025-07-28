"""
Streamlit Cloud Fully Compatible AI Assistant
===========================================

A completely Streamlit Cloud compatible version using:
- Pure Python vector storage (no ChromaDB)
- Sentence transformers for embeddings (with fallback)
- JSON-based persistence
- No native dependencies

Author: Rory Chen
"""

import streamlit as st
import os
import requests
import json
import numpy as np
from typing import List, Dict, Any, Tuple
import uuid
from datetime import datetime
import networkx as nx
from collections import defaultdict
import re
import glob
import pickle
import base64
from pathlib import Path

# Try to import sentence transformers with fallback
try:
    from sentence_transformers import SentenceTransformer
    SENTENCE_TRANSFORMERS_AVAILABLE = True
except:
    SENTENCE_TRANSFORMERS_AVAILABLE = False

# Try to import document processing libraries with fallbacks
try:
    import PyPDF2
    PYPDF2_AVAILABLE = True
except:
    PYPDF2_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except:
    DOCX_AVAILABLE = False

class SimpleVectorStore:
    """Pure Python vector store - fully Streamlit Cloud compatible"""
    
    def __init__(self, persist_path: str = "vector_store.json"):
        self.persist_path = persist_path
        self.documents = []
        self.metadatas = []
        self.ids = []
        self.embeddings = []
        self.conversations = []
        
        # Try to load existing data
        self.load_data()
        
        # Initialize embedder if available
        self.embedder = None
        if SENTENCE_TRANSFORMERS_AVAILABLE:
            try:
                self.embedder = SentenceTransformer('all-MiniLM-L6-v2')
            except:
                pass
    
    def _compute_embedding(self, text: str) -> List[float]:
        """Compute embedding for text"""
        if self.embedder:
            try:
                embedding = self.embedder.encode(text)
                return embedding.tolist()
            except:
                pass
        
        # Fallback: simple TF-IDF-like representation
        words = text.lower().split()
        word_counts = {}
        for word in words:
            if len(word) > 2:  # Skip short words
                word_counts[word] = word_counts.get(word, 0) + 1
        
        # Create a simple vector based on word frequencies
        # This is a very basic fallback but works for simple similarity
        common_words = ['data', 'science', 'machine', 'learning', 'python', 'sql', 'analytics', 
                       'bank', 'insurance', 'axa', 'citic', 'experience', 'project', 'model',
                       'azure', 'cloud', 'tableau', 'dashboard', 'ai', 'deep', 'nlp']
        
        vector = []
        for word in common_words:
            vector.append(word_counts.get(word, 0))
        
        # Normalize
        total = sum(vector) or 1
        return [v / total for v in vector]
    
    def _cosine_similarity(self, vec1: List[float], vec2: List[float]) -> float:
        """Compute cosine similarity between two vectors"""
        try:
            # Convert to numpy arrays
            a = np.array(vec1)
            b = np.array(vec2)
            
            # Compute cosine similarity
            dot_product = np.dot(a, b)
            norm_a = np.linalg.norm(a)
            norm_b = np.linalg.norm(b)
            
            if norm_a == 0 or norm_b == 0:
                return 0.0
            
            return dot_product / (norm_a * norm_b)
        except:
            # Fallback: simple word overlap
            words1 = set(str(vec1).lower().split())
            words2 = set(str(vec2).lower().split())
            intersection = len(words1.intersection(words2))
            union = len(words1.union(words2))
            return intersection / union if union > 0 else 0.0
    
    def add_documents(self, documents: List[str], metadatas: List[Dict], ids: List[str]) -> bool:
        """Add documents to vector store"""
        try:
            for doc, metadata, doc_id in zip(documents, metadatas, ids):
                # Skip if already exists
                if doc_id in self.ids:
                    continue
                
                # Compute embedding
                embedding = self._compute_embedding(doc)
                
                # Add to store
                self.documents.append(doc)
                self.metadatas.append(metadata)
                self.ids.append(doc_id)
                self.embeddings.append(embedding)
            
            # Save data
            self.save_data()
            return True
        except Exception as e:
            st.error(f"Error adding documents: {str(e)}")
            return False
    
    def query_documents(self, query: str, n_results: int = 5) -> Dict:
        """Query documents using vector similarity"""
        if not self.documents:
            return {"documents": [[]], "metadatas": [[]], "distances": [[]]}
        
        try:
            # Compute query embedding
            query_embedding = self._compute_embedding(query)
            
            # Compute similarities
            similarities = []
            for i, doc_embedding in enumerate(self.embeddings):
                similarity = self._cosine_similarity(query_embedding, doc_embedding)
                similarities.append((similarity, i))
            
            # Sort by similarity (descending)
            similarities.sort(reverse=True)
            
            # Get top results
            top_results = similarities[:n_results]
            
            # Format results
            documents = [self.documents[idx] for _, idx in top_results]
            metadatas = [self.metadatas[idx] for _, idx in top_results]
            distances = [1 - sim for sim, _ in top_results]  # Convert similarity to distance
            
            return {
                "documents": [documents],
                "metadatas": [metadatas],
                "distances": [distances]
            }
        
        except Exception as e:
            st.error(f"Error querying documents: {str(e)}")
            return {"documents": [[]], "metadatas": [[]], "distances": [[]]}
    
    def get_collection_count(self) -> int:
        """Get total number of documents"""
        return len(self.documents)
    
    def log_conversation(self, user_query: str, ai_response: str, chunks_used: int = 0) -> bool:
        """Log conversation"""
        try:
            conversation = {
                "id": str(uuid.uuid4()),
                "timestamp": datetime.now().isoformat(),
                "user_query": user_query,
                "ai_response": ai_response,
                "chunks_used": chunks_used
            }
            self.conversations.append(conversation)
            self.save_data()
            return True
        except:
            return False
    
    def get_conversation_count(self) -> int:
        """Get total number of conversations"""
        return len(self.conversations)
    
    def save_data(self):
        """Save data to JSON file"""
        try:
            data = {
                "documents": self.documents,
                "metadatas": self.metadatas,
                "ids": self.ids,
                "embeddings": self.embeddings,
                "conversations": self.conversations
            }
            
            # Use Streamlit session state for persistence in cloud
            st.session_state['vector_store_data'] = data
            
            # Also try to save to file if possible
            try:
                with open(self.persist_path, 'w', encoding='utf-8') as f:
                    json.dump(data, f, ensure_ascii=False, indent=2)
            except:
                pass  # File saving might not work in Streamlit Cloud
                
        except Exception as e:
            st.warning(f"Could not save data: {str(e)}")
    
    def load_data(self):
        """Load data from JSON file or session state"""
        try:
            # Try to load from session state first
            if 'vector_store_data' in st.session_state:
                data = st.session_state['vector_store_data']
                self.documents = data.get('documents', [])
                self.metadatas = data.get('metadatas', [])
                self.ids = data.get('ids', [])
                self.embeddings = data.get('embeddings', [])
                self.conversations = data.get('conversations', [])
                return
            
            # Try to load from file
            if os.path.exists(self.persist_path):
                with open(self.persist_path, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                    self.documents = data.get('documents', [])
                    self.metadatas = data.get('metadatas', [])
                    self.ids = data.get('ids', [])
                    self.embeddings = data.get('embeddings', [])
                    self.conversations = data.get('conversations', [])
        except:
            pass  # Start with empty store if loading fails

class GraphRAGProcessor:
    """Lightweight GraphRAG implementation"""
    
    def __init__(self):
        self.knowledge_graph = nx.Graph()
        self.entity_patterns = {
            'companies': r'(?i)\b(?:China CITIC Bank|CITIC Bank|AXA|Cigna|Ipsos|City University|Education University|CITIC)\b',
            'skills': r'(?i)\b(?:Python|PySpark|SQL|Machine Learning|Deep Learning|NLP|Tableau|Power BI|Azure|AWS|AutoML|MLOps|Data Science|Analytics|AI|Artificial Intelligence)\b',
            'positions': r'(?i)\b(?:AVP|Assistant.*Manager|Data Science Analyst|Research Executive|Manager|Director|Analyst)\b',
            'years': r'\b(?:20\d{2})\b',
            'industries': r'(?i)\b(?:Insurance|Banking|Healthcare|Finance|Market Research|Medicare|ESG|Green|Environmental)\b',
            'projects': r'(?i)\b(?:Go-Green|Go Green|ESG|Dashboard|AutoML|Pipeline|Migration|Campaign)\b',
            'locations': r'(?i)\b(?:Hong Kong|China|Asia)\b'
        }
    
    def extract_entities(self, text: str) -> List[str]:
        """Extract entities from text"""
        entities = []
        for category, pattern in self.entity_patterns.items():
            matches = re.findall(pattern, text, re.IGNORECASE)
            for match in matches:
                entities.append(f"{category}:{match}")
        return entities
    
    def build_graph_from_documents(self, documents: List[str], metadatas: List[Dict]):
        """Build knowledge graph from documents"""
        for i, (doc, metadata) in enumerate(zip(documents, metadatas)):
            entities = self.extract_entities(doc)
            doc_id = f"doc_{i}"
            
            # Add entities as nodes
            for entity in entities:
                if not self.knowledge_graph.has_node(entity):
                    self.knowledge_graph.add_node(entity, documents=[doc_id])
                else:
                    self.knowledge_graph.nodes[entity]['documents'].append(doc_id)
            
            # Connect related entities
            for j, entity1 in enumerate(entities):
                for entity2 in entities[j+1:]:
                    if self.knowledge_graph.has_edge(entity1, entity2):
                        self.knowledge_graph[entity1][entity2]['weight'] += 1
                    else:
                        self.knowledge_graph.add_edge(entity1, entity2, weight=1)
    
    def get_related_entities(self, query: str) -> List[str]:
        """Get entities related to query"""
        query_entities = self.extract_entities(query)
        related = set(query_entities)
        
        for entity in query_entities:
            if entity in self.knowledge_graph:
                neighbors = list(self.knowledge_graph.neighbors(entity))
                related.update(neighbors[:5])
        
        # Also search for partial matches
        query_lower = query.lower()
        for node in self.knowledge_graph.nodes():
            node_value = node.split(':')[-1].lower() if ':' in node else node.lower()
            for word in query_lower.split():
                if len(word) > 2 and word in node_value:
                    related.add(node)
                    neighbors = list(self.knowledge_graph.neighbors(node))
                    related.update(neighbors[:3])
                    break
        
        return list(related)

class QwenAPIClient:
    """Qwen API client"""
    
    def __init__(self, api_key: str = None):
        self.api_key = api_key or os.getenv("QWEN_API_KEY")
        self.api_url = "https://dashscope.aliyuncs.com/api/v1/services/aigc/text-generation/generation"
        self.available = self._test_connection()
    
    def _test_connection(self) -> bool:
        """Test API connection"""
        if not self.api_key:
            return False
        
        try:
            headers = {
                "Authorization": f"Bearer {self.api_key}",
                "Content-Type": "application/json"
            }
            
            data = {
                "model": "qwen-turbo",
                "input": {"prompt": "Hello"},
                "parameters": {"max_tokens": 10, "temperature": 0.1}
            }
            
            response = requests.post(self.api_url, headers=headers, json=data, timeout=10)
            return response.status_code == 200
        except:
            return False
    
    def generate_response(self, query: str, context: str, graph_entities: List[str] = None) -> str:
        """Generate response using Qwen API"""
        if not self.available:
            return "❌ Qwen API is not available. Please contact Rory at chengy823@gmail.com"
        
        system_prompt = """You are Rory Chen's professional AI assistant with access to his comprehensive career information.

CORE INFORMATION ABOUT RORY:
- Current Role: AVP of Data Science at China CITIC Bank International (Nov 2022 - Current)
- 8 years of professional experience (2017-2025)
- Career progression: Research Executive → Data Science Analyst → Assistant Data Science Manager → AVP
- Industries: Market Research, Healthcare (Medicare), Insurance, Banking
- Education: Master's in Quantitative Analysis (2016-2017), Master's in Public Policy (2014-2015), Bachelor's (2011-2014)
- Location: Hong Kong SAR
- Contact: chengy823@gmail.com

TECHNICAL EXPERTISE:
- Programming: Python, Pyspark, SQL (Advanced proficiency)
- Machine Learning: Deep Learning, NLP, Computer Vision, MLOps, AutoML, LangChain
- Cloud Platforms: Azure, Google Cloud, AWS, Databricks, Cloudera CDSW
- Visualization: Tableau, Power BI, Dashboard Development

KEY ACHIEVEMENTS:
- Developed 20+ ML models achieving 1.5x business uplift vs control groups
- Created AutoML pipeline reducing coding effort by 80%
- Designed AI+BI framework with ONE AXA Dashboard
- Led cloud migration from on-premise to Azure infrastructure

MAJOR PROJECTS:
- **Go-Green Home Insurance Product (AXA)**: Led end-to-end data science solution for ESG-related home insurance product. Served as solution partner, engaging with Chief Underwriting Officer and team heads. Worked closely with product, claims, underwriting, and pricing teams to design and develop Hong Kong's first ESG home insurance campaign. Leveraged data science technology to link ESG criteria with risk assessment and pricing models.

Provide professional, accurate responses based on the context."""

        graph_context = ""
        if graph_entities:
            graph_context = f"\n\nGraphRAG Analysis - Related entities: {', '.join(graph_entities[:10])}"
        
        user_prompt = f"""Question: {query}

Context from knowledge base:
{context}
{graph_context}

Please provide a comprehensive, professional response."""

        combined_prompt = f"{system_prompt}\n\n{user_prompt}"
        
        try:
            headers = {
                "Authorization": f"Bearer {self.api_key}",
                "Content-Type": "application/json"
            }
            
            data = {
                "model": "qwen-turbo",
                "input": {"prompt": combined_prompt},
                "parameters": {
                    "max_tokens": 8000,
                    "temperature": 0.7,
                    "top_p": 0.8
                }
            }
            
            response = requests.post(self.api_url, headers=headers, json=data, timeout=45)
            
            if response.status_code == 200:
                result = response.json()
                ai_response = result['output']['text'].strip()
                
                if graph_entities:
                    ai_response += "\n\n*Response enhanced with GraphRAG technology.*"
                
                return ai_response
            else:
                return "❌ Qwen API failed. Please contact Rory at chengy823@gmail.com"
                
        except Exception as e:
            return "❌ Qwen API failed. Please contact Rory at chengy823@gmail.com"

class TextChunker:
    """Text chunking utility"""
    
    def __init__(self, chunk_size: int = 1000, overlap: int = 200, min_chunk_size: int = 100):
        self.chunk_size = chunk_size
        self.overlap = overlap
        self.min_chunk_size = min_chunk_size
    
    def chunk_text(self, text: str, source: str = "unknown") -> List[tuple]:
        """Split text into overlapping chunks"""
        if len(text) <= self.chunk_size:
            return [(text, {"source": source, "chunk_index": 0, "total_chunks": 1})]
        
        chunks = []
        start = 0
        chunk_index = 0
        
        sentences = self._split_into_sentences(text)
        current_chunk = ""
        
        for sentence in sentences:
            if len(current_chunk) + len(sentence) > self.chunk_size and current_chunk:
                if len(current_chunk.strip()) >= self.min_chunk_size:
                    chunks.append((
                        current_chunk.strip(),
                        {
                            "source": source,
                            "chunk_index": chunk_index,
                            "chunk_size": len(current_chunk.strip())
                        }
                    ))
                    chunk_index += 1
                
                if self.overlap > 0 and chunks:
                    overlap_text = self._get_overlap_text(current_chunk, self.overlap)
                    current_chunk = overlap_text + " " + sentence
                else:
                    current_chunk = sentence
            else:
                current_chunk += " " + sentence if current_chunk else sentence
        
        if current_chunk.strip() and len(current_chunk.strip()) >= self.min_chunk_size:
            chunks.append((
                current_chunk.strip(),
                {
                    "source": source,
                    "chunk_index": chunk_index,
                    "chunk_size": len(current_chunk.strip())
                }
            ))
        
        total_chunks = len(chunks)
        for i, (chunk_text, metadata) in enumerate(chunks):
            metadata["total_chunks"] = total_chunks
        
        return chunks
    
    def _split_into_sentences(self, text: str) -> List[str]:
        """Split text into sentences"""
        sentences = re.split(r'(?<=[.!?])\s+', text)
        return [s.strip() for s in sentences if s.strip()]
    
    def _get_overlap_text(self, text: str, overlap_size: int) -> str:
        """Get overlap text from end of chunk"""
        if len(text) <= overlap_size:
            return text
        
        overlap_text = text[-overlap_size:]
        space_index = overlap_text.find(' ')
        if space_index > 0:
            return overlap_text[space_index + 1:]
        return overlap_text

class DocumentProcessor:
    """Process documents with available libraries"""
    
    def __init__(self, chunker: TextChunker = None):
        self.chunker = chunker or TextChunker()
    
    def extract_text_from_pdf(self, pdf_path: str) -> str:
        """Extract text from PDF"""
        if not PYPDF2_AVAILABLE:
            return f"PDF processing not available for {os.path.basename(pdf_path)}"
        
        try:
            import PyPDF2
            text = ""
            with open(pdf_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text.strip():
                        text += page_text + "\n"
            
            return text if len(text.strip()) > 100 else f"No readable text in {os.path.basename(pdf_path)}"
        except Exception as e:
            return f"Error reading {os.path.basename(pdf_path)}: {str(e)}"
    
    def extract_text_from_docx(self, docx_path: str) -> str:
        """Extract text from DOCX"""
        if not DOCX_AVAILABLE:
            return f"DOCX processing not available for {os.path.basename(docx_path)}"
        
        try:
            from docx import Document
            doc = Document(docx_path)
            text = ""
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += cell.text + "\n"
            
            return text if len(text.strip()) > 50 else f"No readable text in {os.path.basename(docx_path)}"
        except Exception as e:
            return f"Error reading {os.path.basename(docx_path)}: {str(e)}"
    
    def process_directory(self, directory: str) -> tuple:
        """Process all documents in directory"""
        raw_documents = []
        raw_metadatas = []
        
        # Process text files
        for ext in ['*.txt', '*.md']:
            files = glob.glob(os.path.join(directory, ext))
            for file_path in files:
                try:
                    with open(file_path, 'r', encoding='utf-8') as f:
                        text = f.read()
                    
                    if text.strip() and len(text) > 100:
                        raw_documents.append(text)
                        raw_metadatas.append({
                            "source": os.path.basename(file_path),
                            "type": "text",
                            "path": file_path
                        })
                except:
                    continue
        
        # Process PDF files
        if PYPDF2_AVAILABLE:
            pdf_files = glob.glob(os.path.join(directory, '*.pdf'))
            for pdf_path in pdf_files:
                text = self.extract_text_from_pdf(pdf_path)
                if text.strip() and len(text) > 100:
                    raw_documents.append(text)
                    raw_metadatas.append({
                        "source": os.path.basename(pdf_path),
                        "type": "pdf",
                        "path": pdf_path
                    })
        
        # Process DOCX files
        if DOCX_AVAILABLE:
            docx_files = glob.glob(os.path.join(directory, '*.docx'))
            for docx_path in docx_files:
                text = self.extract_text_from_docx(docx_path)
                if text.strip() and len(text) > 100:
                    raw_documents.append(text)
                    raw_metadatas.append({
                        "source": os.path.basename(docx_path),
                        "type": "docx",
                        "path": docx_path
                    })
        
        # Apply chunking
        if raw_documents:
            chunked_documents = []
            chunked_metadatas = []
            chunked_ids = []
            
            for doc_index, (document, metadata) in enumerate(zip(raw_documents, raw_metadatas)):
                source = metadata.get("source", f"doc_{doc_index}")
                chunks = self.chunker.chunk_text(document, source)
                
                for chunk_text, chunk_metadata in chunks:
                    combined_metadata = {**metadata, **chunk_metadata}
                    combined_metadata["original_doc_index"] = doc_index
                    
                    chunked_documents.append(chunk_text)
                    chunked_metadatas.append(combined_metadata)
                    chunked_ids.append(f"{source}_chunk_{chunk_metadata['chunk_index']}")
            
            return chunked_documents, chunked_metadatas, chunked_ids
        else:
            return [], [], []

@st.cache_resource
def initialize_system():
    """Initialize the system"""
    try:
        vector_store = SimpleVectorStore()
        graph_rag = GraphRAGProcessor()
        qwen_client = QwenAPIClient()
        doc_processor = DocumentProcessor()
        
        # Check if we need to load documents
        doc_count = vector_store.get_collection_count()
        
        if doc_count == 0:
            # Load default documents
            default_documents = [
                """Rory Chen - Professional Summary
                AVP of Data Science at China CITIC Bank International (Nov 2022 - Current)
                8 years of professional experience in data science and analytics
                Career progression: Research Executive at Ipsos → Data Science Analyst at Cigna → Assistant Data Science Manager at AXA → AVP at CITIC Bank
                Industries: Market Research, Healthcare (Medicare), Insurance, Banking
                Education: Master's in Quantitative Analysis (2016-2017), Master's in Public Policy (2014-2015), Bachelor's (2011-2014)
                Location: Hong Kong SAR
                Contact: chengy823@gmail.com, Phone: 68746551""",
                
                """Rory Chen - Technical Skills and Achievements
                Programming Languages: Python, PySpark, SQL (Advanced proficiency)
                Machine Learning: Deep Learning, NLP, Computer Vision, MLOps, AutoML, LangChain
                Cloud Platforms: Azure, Google Cloud, AWS, Databricks, Cloudera CDSW
                Visualization: Tableau, Power BI, Dashboard Development
                
                Key Achievements:
                - Developed 20+ ML models achieving 1.5x business uplift vs control groups
                - Created AutoML pipeline reducing coding effort by 80%
                - Designed AI+BI framework with ONE AXA Dashboard
                - Led cloud migration from on-premise to Azure infrastructure""",
                
                """Rory Chen - Go-Green Home Insurance Project (AXA)
                Project Name: Go-Green Home Insurance Product
                Department: General Insurance
                
                Process:
                Served as a solution partner, providing an end-to-end data science solution for an ESG-related home insurance product.
                Engaged with the Chief Underwriting Officer and team heads to understand business needs.
                Worked closely with product, claims, underwriting, and pricing teams to design and develop Hong Kong's first ESG home insurance campaign.
                Leveraged data science technology to link ESG criteria with risk assessment and pricing models.
                
                This project represents a significant achievement in combining environmental sustainability with insurance innovation."""
            ]
            
            default_metadatas = [
                {"source": "professional_summary.txt", "type": "text"},
                {"source": "technical_skills.txt", "type": "text"},
                {"source": "go_green_project.txt", "type": "text"}
            ]
            
            default_ids = ["default_1", "default_2", "default_3"]
            
            # Add default documents
            success = vector_store.add_documents(default_documents, default_metadatas, default_ids)
            if success:
                graph_rag.build_graph_from_documents(default_documents, default_metadatas)
                st.success(f"✅ Loaded {len(default_documents)} documents into vector store!")
            
            # Try to process additional documents
            try:
                current_dir = os.path.dirname(os.path.abspath(__file__))
                documents, metadatas, ids = doc_processor.process_directory(current_dir)
                
                if documents:
                    success = vector_store.add_documents(documents, metadatas, ids)
                    if success:
                        graph_rag.build_graph_from_documents(documents, metadatas)
                        doc_sources = list(set([meta.get('source', 'Unknown') for meta in metadatas]))
                        st.success(f"✅ Loaded additional {len(documents)} documents: {', '.join(doc_sources)}")
            except Exception as e:
                st.info(f"Could not process additional documents: {str(e)}")
        
        return vector_store, graph_rag, qwen_client, vector_store.get_collection_count()
    
    except Exception as e:
        st.error(f"Failed to initialize system: {str(e)}")
        return None, None, None, 0

# Streamlit App
st.set_page_config(
    page_title="Rory's AI Assistant - Cloud Compatible",
    page_icon="🤖",
    layout="wide"
)

# Custom CSS
st.markdown("""
<style>
    .main .block-container {
        max-width: 1200px;
        padding-top: 2rem;
    }
    
    .header-container {
        background: linear-gradient(90deg, #1e3c72 0%, #2a5298 100%);
        color: white;
        padding: 2rem;
        border-radius: 10px;
        text-align: center;
        margin-bottom: 2rem;
    }
    
    .status-container {
        background: #f0f2f6;
        padding: 1rem;
        border-radius: 8px;
        border-left: 4px solid #3498db;
        margin-bottom: 2rem;
    }
    
    .info-box {
        background: #e8f4fd;
        padding: 1rem;
        border-radius: 8px;
        border-left: 4px solid #28a745;
        margin-bottom: 1rem;
    }
</style>
""", unsafe_allow_html=True)

# Sidebar
with st.sidebar:
    st.markdown("### 🔑 API Configuration")
    qwen_api_key = st.text_input(
        "Qwen API Key:",
        value="sk-015ea57c8b254c4181d30b2de4259d8b",
        type="password",
        help="Enter your Qwen API key"
    )
    
    if qwen_api_key:
        os.environ["QWEN_API_KEY"] = qwen_api_key
        st.success("✅ API key configured!")
    
    st.markdown("---")
    st.markdown("### 📊 System Info")
    
    st.markdown("**Storage Type:** Pure Python Vector Store")
    st.markdown("**Embeddings:** " + ("Sentence Transformers" if SENTENCE_TRANSFORMERS_AVAILABLE else "TF-IDF Fallback"))
    st.markdown("**PDF Support:** " + ("✅" if PYPDF2_AVAILABLE else "❌"))
    st.markdown("**DOCX Support:** " + ("✅" if DOCX_AVAILABLE else "❌"))
    
    if st.button("🔄 Refresh System"):
        st.cache_resource.clear()
        st.rerun()

# Main header
st.markdown("""
<div class="header-container">
    <h1>🤖 Rory's AI Assistant</h1>
    <p>Streamlit Cloud Compatible • Pure Python Vector Store • GraphRAG • Qwen API</p>
</div>
""", unsafe_allow_html=True)

# Initialize system
with st.spinner("🚀 Initializing AI System..."):
    vector_store, graph_rag, qwen_client, doc_count = initialize_system()

if not vector_store:
    st.error("❌ System initialization failed. Please check the configuration.")
    st.stop()

# System status
api_status = "✅ Connected" if qwen_client.available else "❌ Not Available"
conversation_count = vector_store.get_conversation_count()

st.markdown(f"""
<div class="status-container">
    <strong>🧠 System Status:</strong> {doc_count} knowledge chunks in Pure Python Vector Store | 
    {conversation_count} conversations logged | GraphRAG: ✅ Active | Qwen API: {api_status}
</div>
""", unsafe_allow_html=True)

# Info box about compatibility
st.markdown("""
<div class="info-box">
    <strong>☁️ Streamlit Cloud Compatible:</strong> This version uses pure Python vector storage instead of ChromaDB, 
    making it fully compatible with Streamlit Cloud deployment. All functionality is preserved with intelligent fallbacks.
</div>
""", unsafe_allow_html=True)

# Create layout
col1, col2 = st.columns([1, 2])

with col1:
    st.markdown("""
    ### 👨‍💼 Rory Chen
    **AVP of Data Science**  
    China CITIC Bank International
    
    **Experience:** 8 years (2017-2025)
    
    **Industries:**
    - Banking (China CITIC Bank Int'l)
    - Insurance (AXA)
    - Healthcare (Cigna)
    - Market Research (Ipsos)
                                  
    **Key Skills:**
    - Python, Pyspark, SQL
    - Machine Learning & AI, LangChain
    - Cloud Platforms (Azure, Google Cloud)
    - Data Visualization
    
    **Email:** chengy823@gmail.com
    **Phone:** 68746551
    """)

with col2:
    st.markdown("### 💬 Chat with AI Assistant")
    
    # Initialize chat history
    if "messages" not in st.session_state:
        st.session_state.messages = []
        st.session_state.messages.append({
            "role": "assistant",
            "content": "👋 Hello! I'm Rory's AI assistant powered by GraphRAG and pure Python vector storage. Ask me anything about Rory's professional background!"
        })
    
    # Display chat messages
    for message in st.session_state.messages:
        with st.chat_message(message["role"]):
            st.markdown(message["content"])
    
    # Chat input
    if prompt := st.chat_input("Ask about Rory's experience, skills, or achievements..."):
        # Add user message
        st.session_state.messages.append({"role": "user", "content": prompt})
        
        with st.chat_message("user"):
            st.markdown(prompt)
        
        # Generate response
        with st.chat_message("assistant"):
            with st.spinner("🔍 Processing with GraphRAG..."):
                try:
                    # Query vector store
                    results = vector_store.query_documents(prompt, n_results=5)
                    
                    # Get GraphRAG entities
                    graph_entities = graph_rag.get_related_entities(prompt)
                    
                    # Build context
                    context = ""
                    chunks_used = 0
                    if results['documents'] and results['documents'][0]:
                        context = "\n\n".join(results['documents'][0])
                        chunks_used = len(results['documents'][0])
                    
                    # Generate response
                    response = qwen_client.generate_response(prompt, context, graph_entities)
                    
                    # Display response
                    st.markdown(response)
                    
                    # Display analytics
                    if chunks_used > 0:
                        sources = []
                        if results.get('metadatas') and results['metadatas'][0]:
                            sources = [meta.get('source', 'Unknown') for meta in results['metadatas'][0][:3]]
                        
                        st.markdown(f"""
                        <div style="background: #f0f2f6; padding: 0.5rem; border-radius: 4px; font-size: 0.8rem; margin-top: 0.5rem;">
                            📊 <strong>Query Analysis:</strong> Used {chunks_used} knowledge chunks | 
                            GraphRAG entities: {len(graph_entities)} | 
                            Sources: {', '.join(set(sources)) if sources else 'Default knowledge'}
                        </div>
                        """, unsafe_allow_html=True)
                    
                    # Log conversation
                    vector_store.log_conversation(prompt, response, chunks_used)
                    
                    st.session_state.messages.append({"role": "assistant", "content": response})
                    
                except Exception as e:
                    error_msg = f"Sorry, I encountered an error: {str(e)}"
                    st.error(error_msg)
                    st.session_state.messages.append({"role": "assistant", "content": error_msg})
    
    # Clear chat button
    if st.button("🗑️ Clear Chat"):
        st.session_state.messages = []
        st.rerun()

# Footer
st.markdown("---")
st.markdown("""
<div style="text-align: center; padding: 1rem; color: #666;">
    <p>🤖 <strong>Streamlit Cloud Compatible AI Assistant</strong></p>
    <p>Pure Python Vector Store • GraphRAG • Qwen API</p>
    <p>📧 Contact: chengy823@gmail.com</p>
</div>
""", unsafe_allow_html=True)
