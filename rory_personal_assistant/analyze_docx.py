"""
Simple script to analyze the DOCX content and search for go green project
"""

try:
    from docx import Document
    
    doc = Document('Project experience highlight - Rory.docx')
    text = ''
    
    print("📄 Extracting text from DOCX...")
    
    for paragraph in doc.paragraphs:
        text += paragraph.text + '\n'
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text += cell.text + '\n'
    
    print(f"📊 Total text length: {len(text)} characters")
    
    # Search for the specific project name and related keywords
    keywords = ['Go-Green Home Insurance Product', 'go-green', 'go green', 'ESG', 'green', 'environmental', 'sustainability', 'home insurance', 'AXA', 'homesurance']
    found_keywords = []
    
    text_lower = text.lower()
    
    for keyword in keywords:
        if keyword.lower() in text_lower:
            found_keywords.append(keyword)
            pos = text_lower.find(keyword.lower())
            print(f"\n🔍 Found '{keyword}' at position {pos}")
            start = max(0, pos - 150)
            end = min(len(text), pos + 300)
            context = text[start:end]
            print(f"Context: ...{context}...")
    
    if not found_keywords:
        print("\n❌ No go green/ESG keywords found in the document")
        print("\n📝 Document content preview:")
        print(text[:1000])
        print("\n...")
        print(text[-500:])
    else:
        print(f"\n✅ Found keywords: {found_keywords}")
    
    # Check if this might be the wrong document or if go green is mentioned differently
    if 'AXA' in text:
        print("\n🏢 AXA content found - checking for related projects...")
        axa_pos = text_lower.find('axa')
        axa_context = text[max(0, axa_pos-200):axa_pos+500]
        print(f"AXA context: {axa_context}")

except ImportError:
    print("❌ python-docx library not available")
except FileNotFoundError:
    print("❌ DOCX file not found")
except Exception as e:
    print(f"❌ Error: {e}")
