"""
Simple GraphRAG AI Assistant
============================

A simplified AI assistant using:
- ChromaDB for knowledge base storage
- GraphRAG for intelligent retrieval
- Qwen API for response generation
- Simple, clean frontend design

Author: Rory Chen
"""

import streamlit as st
import os
import chromadb
from chromadb.config import Settings
import requests
import json
from typing import List, Dict, Any
import uuid
from datetime import datetime
import networkx as nx
from collections import defaultdict
import re
import glob
import PyPDF2
import fitz
from docx import Document
from sentence_transformers import SentenceTransformer

class ChromaDBManager:
    """Manages ChromaDB operations for knowledge storage"""
    
    def __init__(self, persist_directory: str = "./chroma_db"):
        self.persist_directory = persist_directory
        self.client = chromadb.PersistentClient(path=persist_directory)
        self.collection_name = "rory_knowledge_base"
        self.conversation_collection_name = "conversation_logs"
        self.collection = self._get_or_create_collection()
        self.conversation_collection = self._get_or_create_conversation_collection()
        
    def _get_or_create_collection(self):
        """Get existing collection or create new one"""
        try:
            return self.client.get_collection(name=self.collection_name)
        except:
            return self.client.create_collection(
                name=self.collection_name,
                metadata={"description": "Rory's professional knowledge base"}
            )
    
    def _get_or_create_conversation_collection(self):
        """Get existing conversation collection or create new one"""
        try:
            return self.client.get_collection(name=self.conversation_collection_name)
        except:
            return self.client.create_collection(
                name=self.conversation_collection_name,
                metadata={"description": "User conversations for model improvement"}
            )
    
    def add_documents(self, documents: List[str], metadatas: List[Dict], ids: List[str]):
        """Add documents to ChromaDB"""
        try:
            self.collection.add(
                documents=documents,
                metadatas=metadatas,
                ids=ids
            )
            return True
        except Exception as e:
            st.error(f"Error adding documents to ChromaDB: {str(e)}")
            return False
    
    def query_documents(self, query: str, n_results: int = 5) -> Dict:
        """Query documents from ChromaDB"""
        try:
            results = self.collection.query(
                query_texts=[query],
                n_results=n_results
            )
            return results
        except Exception as e:
            st.error(f"Error querying ChromaDB: {str(e)}")
            return {"documents": [[]], "metadatas": [[]], "distances": [[]]}
    
    def get_collection_count(self) -> int:
        """Get total number of documents in collection"""
        try:
            return self.collection.count()
        except:
            return 0
    
    def log_conversation(self, user_query: str, ai_response: str, chunks_used: int = 0):
        """Log conversation for model improvement"""
        try:
            conversation_id = str(uuid.uuid4())
            timestamp = datetime.now().isoformat()
            
            conversation_text = f"User: {user_query}\nAI: {ai_response}"
            
            self.conversation_collection.add(
                documents=[conversation_text],
                metadatas=[{
                    "timestamp": timestamp,
                    "user_query": user_query,
                    "ai_response": ai_response,
                    "chunks_used": chunks_used,
                    "type": "conversation"
                }],
                ids=[conversation_id]
            )
            return True
        except Exception as e:
            st.error(f"Error logging conversation: {str(e)}")
            return False
    
    def get_conversation_count(self) -> int:
        """Get total number of conversations logged"""
        try:
            return self.conversation_collection.count()
        except:
            return 0

class GraphRAGProcessor:
    """Simple GraphRAG implementation for enhanced retrieval"""
    
    def __init__(self):
        self.knowledge_graph = nx.Graph()
        self.entity_patterns = {
            'companies': r'\b(?:China CITIC Bank|CITIC Bank|AXA|Cigna|Ipsos|City University|Education University)\b',
            'skills': r'\b(?:Python|R|SQL|Machine Learning|Deep Learning|NLP|Tableau|Power BI|Azure|AWS|AutoML|MLOps)\b',
            'positions': r'\b(?:AVP|Assistant.*Manager|Data Science Analyst|Research Executive|Manager|Director)\b',
            'years': r'\b(?:20\d{2})\b',
            'industries': r'\b(?:Insurance|Banking|Healthcare|Finance|Market Research|Medicare)\b'
        }
    
    def extract_entities(self, text: str) -> List[str]:
        """Extract entities from text"""
        entities = []
        for category, pattern in self.entity_patterns.items():
            matches = re.findall(pattern, text, re.IGNORECASE)
            for match in matches:
                entities.append(f"{category}:{match}")
        return entities
    
    def build_graph_from_documents(self, documents: List[str], metadatas: List[Dict]):
        """Build knowledge graph from documents"""
        for i, (doc, metadata) in enumerate(zip(documents, metadatas)):
            entities = self.extract_entities(doc)
            doc_id = f"doc_{i}"
            
            # Add entities as nodes
            for entity in entities:
                if not self.knowledge_graph.has_node(entity):
                    self.knowledge_graph.add_node(entity, documents=[doc_id])
                else:
                    self.knowledge_graph.nodes[entity]['documents'].append(doc_id)
            
            # Connect related entities
            for j, entity1 in enumerate(entities):
                for entity2 in entities[j+1:]:
                    if self.knowledge_graph.has_edge(entity1, entity2):
                        self.knowledge_graph[entity1][entity2]['weight'] += 1
                    else:
                        self.knowledge_graph.add_edge(entity1, entity2, weight=1)
    
    def get_related_entities(self, query: str) -> List[str]:
        """Get entities related to query"""
        query_entities = self.extract_entities(query)
        related = set(query_entities)
        
        for entity in query_entities:
            if entity in self.knowledge_graph:
                neighbors = list(self.knowledge_graph.neighbors(entity))
                related.update(neighbors[:5])  # Limit to top 5 neighbors
        
        return list(related)

class QwenAPIClient:
    """Qwen API client for response generation"""
    
    def __init__(self, api_key: str = None):
        self.api_key = api_key or os.getenv("QWEN_API_KEY")
        self.api_url = "https://dashscope.aliyuncs.com/api/v1/services/aigc/text-generation/generation"
        self.available = self._test_connection()
    
    def _test_connection(self) -> bool:
        """Test API connection"""
        if not self.api_key:
            return False
        
        try:
            headers = {
                "Authorization": f"Bearer {self.api_key}",
                "Content-Type": "application/json"
            }
            
            data = {
                "model": "qwen-turbo",
                "input": {"prompt": "Hello"},
                "parameters": {"max_tokens": 10, "temperature": 0.1}
            }
            
            response = requests.post(self.api_url, headers=headers, json=data, timeout=10)
            return response.status_code == 200
        except:
            return False
    
    def generate_response(self, query: str, context: str, graph_entities: List[str] = None) -> str:
        """Generate response using Qwen API only - no fallbacks"""
        if not self.available:
            return "❌ Qwen API is not available. Please contact Rory at chengy823@gmail.com"
        
        # Build comprehensive prompt
        system_prompt = """You are Rory Chen's professional AI assistant with access to his comprehensive career information through GraphRAG technology.

CORE INFORMATION ABOUT RORY:
- Current Role: AVP of Data Science at China CITIC Bank International (Nov 2022 - Current)
- 8 years of professional experience (2017-2025)
- Career progression: Research Executive → Data Science Analyst → Assistant Data Science Manager → AVP
- Industries: Market Research, Healthcare (Medicare), Insurance, Banking
- Education: Master's in Quantitative Analysis (2016-2017), Master's in Public Policy (2014-2015), Bachelor's (2011-2014)
- Location: Hong Kong SAR
- Contact: chengy823@gmail.com

TECHNICAL EXPERTISE:
- Programming: Python, Pyspark, SQL (Advanced proficiency)
- Machine Learning: Deep Learning, NLP, Computer Vision, MLOps, AutoML, LangChain
- Data Processing and ETL
- Pega Customer Decision Hub
- Analytics: Predictive Modeling, Customer Analytics, Time Series Analysis
- Cloud Platforms: Azure, Google Cloud, Databricks, Cloudera CDSW
- Visualization: Tableau, Power BI, Dashboard Development

KEY ACHIEVEMENTS:
- Developed 20+ ML models achieving 1.5x business uplift vs control groups
- Created AutoML pipeline reducing coding effort by 80%
- Designed AI+BI framework with ONE AXA Dashboard 
- Led cloud migration from on-premise to Azure infrastructure

CAREER PROGRESSION REASONING:
- Ipsos (Market Research) → Cigna (Healthcare Analytics) → AXA (Insurance) → CITIC Bank (Banking)
- Client relationships at Ipsos likely included insurance/banking companies, facilitating career transitions
- Progressive technical leadership from individual contributor to AVP level
- Cross-industry expertise demonstrates adaptability and broad business acumen

Provide professional, accurate responses based on the context and demonstrate reasoning."""

        # Add GraphRAG context if available
        graph_context = ""
        if graph_entities:
            graph_context = f"\n\nGraphRAG Analysis - Related entities: {', '.join(graph_entities[:10])}"
        
        user_prompt = f"""Question: {query}

Context from knowledge base:
{context}
{graph_context}

Please provide a comprehensive, professional response."""

        combined_prompt = f"{system_prompt}\n\n{user_prompt}"
        
        # Single attempt with longer timeout - Qwen API only
        try:
            headers = {
                "Authorization": f"Bearer {self.api_key}",
                "Content-Type": "application/json"
            }
            
            data = {
                "model": "qwen-turbo",
                "input": {"prompt": combined_prompt},
                "parameters": {
                    "max_tokens": 5000,
                    "temperature": 0.7,
                    "top_p": 0.8
                }
            }
            
            response = requests.post(self.api_url, headers=headers, json=data, timeout=45)
            
            if response.status_code == 200:
                result = response.json()
                ai_response = result['output']['text'].strip()
                
                if graph_entities:
                    ai_response += "\n\n*Response enhanced with GraphRAG technology.*"
                
                return ai_response
            else:
                return "❌ Qwen API failed. Please contact Rory at chengy823@gmail.com"
                
        except Exception as e:
            return "❌ Qwen API failed. Please contact Rory at chengy823@gmail.com"

class DocumentProcessor:
    """Process documents for knowledge base"""
    
    def extract_text_from_pdf(self, pdf_path: str) -> str:
        """Extract text from PDF with OCR fallback"""
        text = ""
        
        # Try PyMuPDF first
        try:
            doc = fitz.open(pdf_path)
            for page in doc:
                page_text = page.get_text()
                if page_text.strip():
                    text += page_text + "\n"
            doc.close()
            
            if len(text.strip()) > 100:
                return text
        except:
            pass
        
        # Fallback to PyPDF2
        try:
            with open(pdf_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text.strip():
                        text += page_text + "\n"
            
            if len(text.strip()) > 100:
                return text
        except:
            pass
        
        return f"Could not extract text from {os.path.basename(pdf_path)}"
    
    def extract_text_from_docx(self, docx_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            doc = Document(docx_path)
            text = ""
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += cell.text + "\n"
            
            return text if len(text.strip()) > 50 else f"No readable text in {os.path.basename(docx_path)}"
        except Exception as e:
            return f"Error reading {os.path.basename(docx_path)}: {str(e)}"
    
    def process_directory(self, directory: str) -> tuple:
        """Process all documents in directory"""
        documents = []
        metadatas = []
        ids = []
        
        # Process text files
        for ext in ['*.txt', '*.md']:
            files = glob.glob(os.path.join(directory, ext))
            for file_path in files:
                try:
                    with open(file_path, 'r', encoding='utf-8') as f:
                        text = f.read()
                    
                    if text.strip() and len(text) > 100:
                        documents.append(text)
                        metadatas.append({
                            "source": os.path.basename(file_path),
                            "type": "text",
                            "path": file_path
                        })
                        ids.append(f"text_{len(ids)}")
                except:
                    continue
        
        # Process PDF files
        pdf_files = glob.glob(os.path.join(directory, '*.pdf'))
        for pdf_path in pdf_files:
            text = self.extract_text_from_pdf(pdf_path)
            if text.strip() and len(text) > 100:
                documents.append(text)
                metadatas.append({
                    "source": os.path.basename(pdf_path),
                    "type": "pdf",
                    "path": pdf_path
                })
                ids.append(f"pdf_{len(ids)}")
        
        # Process DOCX files
        docx_files = glob.glob(os.path.join(directory, '*.docx'))
        for docx_path in docx_files:
            text = self.extract_text_from_docx(docx_path)
            if text.strip() and len(text) > 100:
                documents.append(text)
                metadatas.append({
                    "source": os.path.basename(docx_path),
                    "type": "docx",
                    "path": docx_path
                })
                ids.append(f"docx_{len(ids)}")
        
        return documents, metadatas, ids

@st.cache_resource
def initialize_system():
    """Initialize the complete system"""
    try:
        # Initialize components
        chroma_db = ChromaDBManager()
        graph_rag = GraphRAGProcessor()
        qwen_client = QwenAPIClient()
        doc_processor = DocumentProcessor()
        
        # Check if knowledge base needs to be populated
        doc_count = chroma_db.get_collection_count()
        
        if doc_count == 0:
            # Process documents from current directory
            current_dir = os.path.dirname(os.path.abspath(__file__))
            documents, metadatas, ids = doc_processor.process_directory(current_dir)
            
            if documents:
                # Add to ChromaDB
                success = chroma_db.add_documents(documents, metadatas, ids)
                if success:
                    # Build GraphRAG
                    graph_rag.build_graph_from_documents(documents, metadatas)
                    st.success(f"✅ Loaded {len(documents)} documents into knowledge base!")
                else:
                    st.error("❌ Failed to load documents into ChromaDB")
            else:
                st.warning("⚠️ No documents found to load")
        else:
            # Load existing documents for GraphRAG
            try:
                all_docs = chroma_db.collection.get()
                if all_docs['documents']:
                    graph_rag.build_graph_from_documents(all_docs['documents'], all_docs['metadatas'])
            except:
                pass
        
        return chroma_db, graph_rag, qwen_client, chroma_db.get_collection_count()
    
    except Exception as e:
        st.error(f"Failed to initialize system: {str(e)}")
        return None, None, None, 0

# Streamlit App Configuration
st.set_page_config(
    page_title="Rory's AI Assistant",
    page_icon="🤖",
    layout="wide",
    initial_sidebar_state="collapsed"
)

# Custom CSS for dark theme design
st.markdown("""
<style>
    .main .block-container {
        font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
        max-width: 1200px;
        padding-top: 2rem;
        background-color: #0e1117;
    }
    
    .header-container {
        background: linear-gradient(90deg, #1e3c72 0%, #2a5298 100%);
        color: white;
        padding: 2rem;
        border-radius: 10px;
        text-align: center;
        margin-bottom: 2rem;
        box-shadow: 0 4px 6px rgba(0,0,0,0.3);
    }
    
    .status-container {
        background: #262730;
        color: #ffffff;
        padding: 1rem;
        border-radius: 8px;
        border-left: 4px solid #3498db;
        margin-bottom: 2rem;
        box-shadow: 0 2px 4px rgba(0,0,0,0.3);
    }
    
    .chat-container {
        background: #262730;
        color: #ffffff;
        border-radius: 10px;
        padding: 1rem;
        box-shadow: 0 4px 6px rgba(0,0,0,0.3);
        border: 1px solid #404040;
    }
    
    .info-box {
        background: #1e1e1e;
        color: #ffffff;
        padding: 1rem;
        border-radius: 8px;
        border-left: 4px solid #28a745;
        margin-bottom: 1rem;
        box-shadow: 0 2px 4px rgba(0,0,0,0.3);
    }
    
    .privacy-notice {
        background: #2d3748;
        color: #e2e8f0;
        padding: 0.8rem;
        border-radius: 6px;
        border-left: 3px solid #f6ad55;
        margin-bottom: 1rem;
        font-size: 0.9rem;
        box-shadow: 0 2px 4px rgba(0,0,0,0.2);
    }
    
    .chunk-info {
        background: #1a202c;
        color: #cbd5e0;
        padding: 0.5rem;
        border-radius: 4px;
        font-size: 0.8rem;
        margin-top: 0.5rem;
        border: 1px solid #4a5568;
    }
    
    /* Dark theme for sidebar */
    .css-1d391kg {
        background-color: #262730;
    }
    
    /* Dark theme for text inputs */
    .stTextInput > div > div > input {
        background-color: #1e1e1e;
        color: #ffffff;
        border: 1px solid #404040;
    }
    
    /* Dark theme for chat messages */
    .stChatMessage {
        background-color: #2d3748;
        border: 1px solid #404040;
    }
</style>
""", unsafe_allow_html=True)

# Sidebar for API configuration
with st.sidebar:
    st.markdown("### 🔑 API Configuration")
    qwen_api_key = st.text_input(
        "Qwen API Key:",
        value="sk-a53be657f784477289f546d32fc2cfb4",
        type="password",
        help="Enter your Qwen API key"
    )
    
    if qwen_api_key:
        os.environ["QWEN_API_KEY"] = qwen_api_key
        st.success("✅ API key configured!")
    
    st.markdown("---")
    st.markdown("### 📊 System Info")
    
    if st.button("🔄 Refresh Knowledge Base"):
        st.cache_resource.clear()
        st.rerun()

# Main header
st.markdown("""
<div class="header-container">
    <h1>🤖 Rory's GraphRAG AI Assistant</h1>
    <p>ChromaDB • GraphRAG • Qwen API</p>
</div>
""", unsafe_allow_html=True)

# Initialize system
with st.spinner("🚀 Initializing GraphRAG AI System..."):
    chroma_db, graph_rag, qwen_client, doc_count = initialize_system()

if not chroma_db:
    st.error("❌ System initialization failed. Please check the configuration.")
    st.stop()

# System status with conversation count
api_status = "✅ Connected" if qwen_client.available else "❌ Not Available"
conversation_count = chroma_db.get_conversation_count()
st.markdown(f"""
<div class="status-container">
    <strong>🧠 System Status:</strong> {doc_count} knowledge chunks in ChromaDB | 
    {conversation_count} conversations logged | GraphRAG: ✅ Active | Qwen API: {api_status}
</div>
""", unsafe_allow_html=True)

# Privacy Notice
st.markdown("""
<div class="privacy-notice">
    <strong>📋 Privacy Notice:</strong> Your questions and our responses are stored in our database for model improvement and learning purposes. This helps us provide better assistance over time. By using this system, you consent to this data collection.
</div>
""", unsafe_allow_html=True)

# Create layout
col1, col2 = st.columns([1, 2])

with col1:
    st.markdown("""
    ### 👨‍💼 Rory Chen
    **AVP of Data Science**  
    China CITIC Bank International
    
    **Experience:** 8 years (2017-2025)
    
    **Industries:**
    - Banking (China CITIC Bank Int'l)
    - Insurance (AXA)
    - Insurance (Cigna)
    - Market Research (Ipsos)
                                  
    **Key Skills:**
    - Python, Pyspark, SQL
    - Machine Learning & AI, LangChain
    - Cloud Platforms (Azure, Google Cloud)
    - Pega Customer Decision Hub
    - Data Visualization
    
    **Email:** chengy823@gmail.com
    **Phone:** 68746551
    """)

with col2:
    st.markdown('<div class="chat-container">', unsafe_allow_html=True)
    st.markdown("### 💬 Chat with AI Assistant")
    
    # Initialize chat history
    if "messages" not in st.session_state:
        st.session_state.messages = []
        st.session_state.messages.append({
            "role": "assistant",
            "content": "👋 Hello! I'm Rory's AI assistant powered by GraphRAG and Qwen API. Ask me anything about Rory's professional background!"
        })
    
    # Display chat messages
    for message in st.session_state.messages:
        with st.chat_message(message["role"]):
            st.markdown(message["content"])
    
    # Chat input
    if prompt := st.chat_input("Ask about Rory's experience, skills, or achievements..."):
        # Add user message
        st.session_state.messages.append({"role": "user", "content": prompt})
        
        with st.chat_message("user"):
            st.markdown(prompt)
        
        # Generate response
        with st.chat_message("assistant"):
            with st.spinner("🔍 Processing with GraphRAG..."):
                try:
                    # Query ChromaDB
                    chroma_results = chroma_db.query_documents(prompt, n_results=3)
                    
                    # Get GraphRAG entities
                    graph_entities = graph_rag.get_related_entities(prompt)
                    
                    # Build context
                    context = ""
                    chunks_used = 0
                    if chroma_results['documents'] and chroma_results['documents'][0]:
                        context = "\n\n".join(chroma_results['documents'][0])
                        chunks_used = len(chroma_results['documents'][0])
                    
                    # Generate response
                    response = qwen_client.generate_response(prompt, context, graph_entities)
                    
                    # Display response
                    st.markdown(response)
                    
                    # Display chunk information
                    if chunks_used > 0:
                        st.markdown(f"""
                        <div class="chunk-info">
                            📊 <strong>Query Analysis:</strong> Used {chunks_used} knowledge chunks | 
                            GraphRAG entities: {len(graph_entities)} | 
                            Sources: {', '.join(set([meta.get('source', 'Unknown') for meta in chroma_results.get('metadatas', [[]])[0][:3]]))}
                        </div>
                        """, unsafe_allow_html=True)
                    
                    # Log conversation to database
                    chroma_db.log_conversation(prompt, response, chunks_used)
                    
                    st.session_state.messages.append({"role": "assistant", "content": response})
                    
                except Exception as e:
                    error_msg = f"Sorry, I encountered an error: {str(e)}"
                    st.error(error_msg)
                    st.session_state.messages.append({"role": "assistant", "content": error_msg})
    
    # Clear chat button
    if st.button("🗑️ Clear Chat"):
        st.session_state.messages = []
        st.rerun()
    
    st.markdown('</div>', unsafe_allow_html=True)

# Footer
st.markdown("---")
st.markdown("""
<div style="text-align: center; padding: 1rem; color: #666;">
    <p>🤖 <strong>Simple GraphRAG AI Assistant</strong></p>
    <p>ChromaDB • GraphRAG • Qwen API</p>
    <p>📧 Contact: chengy823@gmail.com</p>
</div>
""", unsafe_allow_html=True)
